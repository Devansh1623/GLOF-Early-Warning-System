import os
import base64
import json
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_mermaid_image(mermaid_code):
    state = {
        "code": mermaid_code,
        "mermaid": {
            "theme": "default"
        }
    }
    b64 = base64.urlsafe_b64encode(json.dumps(state).encode('utf-8')).decode('ascii')
    url = f"https://mermaid.ink/img/{b64}?type=png"
    
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to fetch image: {e}")
        return None

def create_week5_report():
    doc = Document()
    
    # Title
    title = doc.add_heading("Week 5 Project Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run("Project Title: ").bold = True
    p.add_run("GLOF Early Warning System (GLOFWatch)\n")
    p.add_run("Phase 3: ").bold = True
    p.add_run("Object-Oriented Design\n")
    p.add_run("Focus: ").bold = True
    p.add_run("Defining the 'How' - architecture, patterns, and component interactions.\n")
    p.add_run("Date Range: ").bold = True
    p.add_run("16 - 21 February 2026\n")
    
    # 1. Executive Summary
    doc.add_heading("1. Executive Summary: Interaction Modeling", level=1)
    doc.add_paragraph(
        "During Week 5, the project advanced into Phase 3: Object-Oriented Design. We transitioned from "
        "analyzing system requirements (the 'What') to defining the operational architecture (the 'How'). "
        "The primary focus was Interaction Modeling, ensuring that our decentralized components (Flask, MongoDB, Redis, and Celery) "
        "communicate efficiently without bottlenecks."
    )
    
    # 2. Defining Component Interactions
    doc.add_heading("2. Defining Component Interactions", level=1)
    doc.add_paragraph(
        "Based on our architecture, the primary critical path is the ingestion of telemetry data that triggers a high-risk alert. "
        "The interaction required involves validation layers, persistent data storage, algorithmic inference, publish/subscribe routing, "
        "and background job queuing. To maintain a strict separation of concerns, the interactions were modeled linearly over time to "
        "highlight synchronization points and asynchronous hand-offs."
    )
    
    # 3. UML Deliverable: Sequence Diagram
    doc.add_heading("3. UML Deliverable: Sequence Diagram", level=1)
    doc.add_paragraph(
        "The following UML Sequence Diagram maps the exact life cycle of an incoming telemetry request from an IoT sensor. "
        "It details the synchronous interactions with the Marshmallow validation schemas and the ML Risk Engine, alongside the "
        "asynchronous triggers utilized for Redis (SSE) and Celery (Email) when an anomaly is detected."
    )
    
    sequence_diagram_code = """
sequenceDiagram
    autonumber
    actor Sensor as IoT Sensor
    participant API as Flask API (/api/telemetry)
    participant schema as Marshmallow Validator
    participant DB as MongoDB
    participant Risk as Risk Engine (XGBoost/Costa)
    participant Redis as Redis Pub/Sub
    participant Celery as Celery Worker
    actor Admin as React Dashboard (SSE)

    Sensor->>API: POST /api/telemetry (Auth: API_KEY)
    API->>schema: Validate JSON Payload
    schema-->>API: Validated Data
    
    par Async persistence
        API->>DB: Insert Telemetry Reading
        API->>Risk: calculate_risk(lake_id, metrics)
    end
    
    Risk-->>API: Returns Tier (e.g., CRITICAL)
    
    alt If Risk Tier == High or Critical
        API->>DB: Insert Alert Record
        
        par Trigger Notifications
            API->>Redis: PUBLISH event to 'alerts' channel
            API->>Celery: queue_email_task(alert_details)
        end
        
        Redis-->>Admin: Server-Sent Event (SSE) Push
        Celery-->>Admin: Dispatch Email (via Resend API)
    end
    
    API-->>Sensor: 200 OK Response
"""
    print("Fetching Sequence Diagram...")
    img = get_mermaid_image(sequence_diagram_code)
    if img:
        doc.add_picture(img, width=Inches(6.5))
    else:
        doc.add_paragraph("[Failed to load Sequence Diagram]", style="Intense Quote")

    # 4. Plan for Next Week
    doc.add_heading("4. Plan for Next Week (Week 6)", level=1)
    doc.add_paragraph("Entering Week 6, the project will continue within the Object-Oriented Design phase. The primary objectives include:", style='List Bullet')
    doc.add_paragraph("Detailed Class Design & Packaging: Applying SOLID principles and formalized Design Patterns to the system architecture.", style='List Bullet')
    doc.add_paragraph("Develop a Detailed Class Diagram explicitly outlining methods, attributes, visibility modifiers (+, -, #), and functional interfaces.", style='List Bullet')
    doc.add_paragraph("Create a Package Diagram to organize related classes into logical modules and bounded subsystems (e.g., core, routes, db).", style='List Bullet')

    # Footer
    p_footer = doc.add_paragraph("\n\nMentor Signature: __________________      Project Guide Signature: __________________")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(r"c:\GLOF\reports\in_depth", exist_ok=True)
    filename = r"c:\GLOF\reports\in_depth\GLOF_Week_5_Report_InDepth.docx"
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    create_week5_report()
