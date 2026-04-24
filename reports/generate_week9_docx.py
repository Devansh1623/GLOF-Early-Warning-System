import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_week9_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('GLOF Early Warning System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Minor Project Report - Week 9')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Phase 4: Implementation & CI/CD Integration').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Focus: CI/CD Pipeline Setup').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: 30 March - 04 April 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_paragraph(doc, 'During Week 9, the project engineering efforts focused entirely on establishing robust Continuous Integration and Continuous Deployment (CI/CD) pipelines. In keeping with modern DevOps practices, we successfully integrated automated build and test pipelines to enforce code quality, ensure functional integrity, and accelerate feature merging without regressions.')
    
    # 2. CI/CD Architecture
    add_heading(doc, '2. Pipeline Configuration via GitHub Actions', 1)
    add_paragraph(doc, 'GitHub Actions was chosen as our CI/CD provider due to its tight integration with our existing repository ecosystem. We implemented an automated workflow file (.github/workflows/ci.yml) configured to trigger on every push to the `main` and `develop` branches, as well as on any Pull Requests targeting production.')
    
    add_paragraph(doc, 'The pipeline is organized into three distinct, parallelized jobs that mimic the full-stack architecture:', bold=True)
    
    add_paragraph(doc, 'Job 1: Backend Testing & Linting')
    p_steps1 = doc.add_paragraph(style='List Bullet')
    p_steps1.add_run('Environment: ').bold = True
    p_steps1.add_run('Ubuntu runners with Python 3.11 installed via setup-python caching.')
    p_steps1 = doc.add_paragraph(style='List Bullet')
    p_steps1.add_run('Linting: ').bold = True
    p_steps1.add_run('Flake8 is executed to enforce PEP-8 standards (max-line-length=120) across all Flask endpoints.')
    p_steps1 = doc.add_paragraph(style='List Bullet')
    p_steps1.add_run('Unit Testing: ').bold = True
    p_steps1.add_run('Pytest calculates logic correctness and ensures coverage bounds (--cov-fail-under=70) on core algorithms and routing.')

    add_paragraph(doc, '\nJob 2: Frontend Build Validation')
    p_steps2 = doc.add_paragraph(style='List Bullet')
    p_steps2.add_run('Environment: ').bold = True
    p_steps2.add_run('Node.js version 20 via setup-node caching.')
    p_steps2 = doc.add_paragraph(style='List Bullet')
    p_steps2.add_run('Build Process: ').bold = True
    p_steps2.add_run('Executes a complete React build (`npm run build`). This immediately catches any syntax, JSX, or dependency import errors before they can reach the production CDN pipeline.')

    add_paragraph(doc, '\nJob 3: Infrastructure / Docker Compose Build')
    p_steps3 = doc.add_paragraph(style='List Bullet')
    p_steps3.add_run('Dependency: ').bold = True
    p_steps3.add_run('This job depends on the success of both the Frontend and Backend pipelines.')
    p_steps3 = doc.add_paragraph(style='List Bullet')
    p_steps3.add_run('Docker Builds: ').bold = True
    p_steps3.add_run('Runs `docker compose build` to ensure the containerization recipes (local Dockerfiles) evaluate perfectly across container images. Validates structural readiness for deployment to Render.com platform.')

    # 3. Benefits Achieved
    add_heading(doc, '3. DevOps Achievements', 1)
    add_paragraph(doc, 'By configuring automated build triggers on every `git push`, the application lifecycle now includes an impenetrable safety net. Developers cannot merge broken code into the `main` branch. Automated pipeline tests save human review time, prevent system outages upon deployment, and form a fundamental prerequisite for agile continuous delivery.')

    # 4. Plan for Next Week
    add_heading(doc, '4. Plan for Next Week (Week 10)', 1)
    doc.add_paragraph('Entering Week 10, the project will wrap up the Implementation & CI/CD phase. The primary objectives are:', style='List Bullet')
    doc.add_paragraph('Implement Sprint 2 Core features, tying the React UI comprehensively with endpoints via custom hooks.', style='List Bullet')
    doc.add_paragraph('Execute final holistic E2E (End to End) application testing before transitioning phases.', style='List Bullet')

    # Ensure output directory exists before saving
    output_dir = r"c:\GLOF\reports\in_depth"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "GLOF_Week_9_Report_InDepth.docx")
    
    try:
        doc.save(file_path)
        print(f"Successfully generated {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_week9_report()
