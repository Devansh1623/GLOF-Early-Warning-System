import os
import base64
import json
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_mermaid_image(mermaid_code):
    """Fetches a PNG from mermaid.ink given a string of Mermaid syntax."""
    state = {
        "code": mermaid_code,
        "mermaid": {
            "theme": "default"
        }
    }
    b64 = base64.urlsafe_b64encode(json.dumps(state).encode('utf-8')).decode('ascii')
    url = f"https://mermaid.ink/img/{b64}?type=png"
    
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to fetch image: {e}")
        return None

def create_week2_report():
    doc = Document()
    
    # Title
    title = doc.add_heading("Week 2 Project Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run("Project Title: ").bold = True
    p.add_run("GLOF Early Warning System (GLOFWatch)\n")
    p.add_run("Phase 1: ").bold = True
    p.add_run("Planning & Requirement Analysis\n")
    p.add_run("Focus: ").bold = True
    p.add_run("Functional Modeling: Actors, Use Cases, and System Boundaries.\n")
    p.add_run("Date Range: ").bold = True
    p.add_run("26 - 31 January 2026\n")
    
    # 1. Executive Summary & Functional Modeling
    doc.add_heading("1. Executive Summary & Functional Modeling", level=1)
    doc.add_paragraph(
        "During Week 2, the primary objective was Functional Modeling. The team transitioned from abstract "
        "requirements defined in Week 1 to specific, actionable system behaviors. We identified all primary "
        "and secondary actors interacting with the GLOF Early Warning System and mapped their interactions "
        "with specific Use Cases."
    )
    
    # 2. Actors & System Boundaries
    doc.add_heading("2. Identifying Actors and System Boundaries", level=1)
    p_actors = doc.add_paragraph(
        "Based on our architecture leveraging Flask and IoT telemetry arrays, four distinct actors were identified:"
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Actor'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description / Role'
    
    actor_items = [
        ("IoT Sensor Node", "Primary (Automated)", "Remote stations pushing telemetry (rainfall, temperature, area change) to the /api/telemetry REST API."),
        ("Admin / Expert", "Primary (Human)", "Glaciologists or system admins accessing the React dashboard to monitor the lakes, add thresholds, and manage users."),
        ("ML Risk Engine", "Secondary (Automated)", "The internal XGBoost & Isolation Forest inference engine evaluating ingested telemetry against historical datasets."),
        ("Public / Subscribed User", "Secondary (Human)", "Downstream residents and local authorities receiving automated SMS/email alerts via the Celery Worker queue.")
    ]
    
    for actor, ttype, desc in actor_items:
        row_cells = table.add_row().cells
        row_cells[0].text = actor
        row_cells[1].text = ttype
        row_cells[2].text = desc

    # 3. UML Deliverable: Use Case Diagram
    doc.add_heading("3. UML Deliverable: Use Case Diagram", level=1)
    doc.add_paragraph(
        "The following diagram defines the system boundaries of the GLOF Early Warning System. It clearly illustrates "
        "the relationship between the actors and the core modules (Telemetry Ingestion, Risk Assessment, Alert Dispatch, and UI)."
    )
    
    mermaid_code = """
flowchart TB
classDef actorStyle fill:#fff,stroke:#333,stroke-width:2px;
classDef usecaseStyle fill:#e1f5fe,stroke:#0288d1,stroke-width:2px;

subgraph "GLOF Early Warning System"
    UC1([Ingest Telemetry Data]):::usecaseStyle
    UC2([Calculate Risk Score]):::usecaseStyle
    UC2a([Calculate Costa Formula]):::usecaseStyle
    UC2b([XGBoost Anomaly Inference]):::usecaseStyle
    UC3([Trigger Alerts via SSE / Email]):::usecaseStyle
    UC4([View Analytics Dashboard]):::usecaseStyle
    UC5([Manage Lake Thresholds]):::usecaseStyle

    UC1 -. "<<includes>>" .-> UC2
    UC2 -. "<<includes>>" .-> UC2a
    UC2 -. "<<includes>>" .-> UC2b
    UC2 --> UC3
end

Sensor((IoT Sensor Node)):::actorStyle
Admin((Admin / Expert)):::actorStyle
ML((Machine Learning Model)):::actorStyle

Sensor --> UC1
ML --> UC2b
Admin --> UC4
Admin --> UC5
"""

    print("Fetching Use Case Diagram Image...")
    img_buffer = get_mermaid_image(mermaid_code)
    if img_buffer:
        doc.add_picture(img_buffer, width=Inches(6.0))
    else:
        doc.add_paragraph("[Failed to load Mermaid Diagram]", style="Intense Quote")

    # 4. Use Case Descriptions
    doc.add_heading("4. Use Case Descriptions (Primary Flows)", level=1)
    
    # UC1
    doc.add_heading("UC-01: Ingest Telemetry Data", level=2)
    uc1_p = doc.add_paragraph(style='List Bullet')
    uc1_p.add_run("Primary Actor: ").bold = True; uc1_p.add_run("IoT Sensor Node")
    uc1_p = doc.add_paragraph(style='List Bullet')
    uc1_p.add_run("Pre-condition: ").bold = True; uc1_p.add_run("Sensor must have a valid SENSOR_API_KEY generated via HMAC-SHA256.")
    uc1_p = doc.add_paragraph(style='List Bullet')
    uc1_p.add_run("Main Success Scenario: ").bold = True
    uc1_p.add_run("The sensor posts a JSON payload to the Flask backend. Schema validation (Marshmallow) ensures temperature, rainfall, and volume metrics are present. Data is persisted to MongoDB with a TTL index.")
    
    # UC2
    doc.add_heading("UC-02: Calculate Risk Score", level=2)
    uc2_p = doc.add_paragraph(style='List Bullet')
    uc2_p.add_run("Primary Actor: ").bold = True; uc2_p.add_run("ML Risk Engine")
    uc2_p = doc.add_paragraph(style='List Bullet')
    uc2_p.add_run("Included Use Cases: ").bold = True; uc2_p.add_run("Costa Formula Evaluation, XGBoost Inference.")
    uc2_p = doc.add_paragraph(style='List Bullet')
    uc2_p.add_run("Main Success Scenario: ").bold = True
    uc2_p.add_run("Triggered immediately after UC-01. The system computes the base vulnerability using real-time Lake depth metrics. It then passes the data to the trained XGBoost model to detect anomalies. Returns a final risk tier (Low, Moderate, High, Critical).")
    
    # UC3
    doc.add_heading("UC-03: Trigger Alerts", level=2)
    uc3_p = doc.add_paragraph(style='List Bullet')
    uc3_p.add_run("Primary Actor: ").bold = True; uc3_p.add_run("System / Public User")
    uc3_p = doc.add_paragraph(style='List Bullet')
    uc3_p.add_run("Pre-condition: ").bold = True; uc3_p.add_run("Risk score classified as High or Critical.")
    uc3_p = doc.add_paragraph(style='List Bullet')
    uc3_p.add_run("Main Success Scenario: ").bold = True
    uc3_p.add_run("A Redis Pub/Sub message is broadcast to all active React Dashboard clients via SSE. Simultaneously, a Celery asynchronous email task is dispatched via the Resend API to the Lake's subscriber list.")

    # 5. Tooling: Jira Mapping
    doc.add_heading("5. Tooling: Mapping Use Cases to Jira User Stories", level=1)
    doc.add_paragraph(
        "To ensure development tracks closely with the functional requirements, the finalized Use Cases were mapped "
        "directly onto the Jira Product Backlog created in Week 1."
    )
    p_jira = doc.add_paragraph(style='List Bullet')
    p_jira.add_run("UC-01 (Ingestion) -> Epic: Data Pipeline -> ").bold = True
    p_jira.add_run("Story GLOF-01: Build REST endpoint with Marshmallow validation.")
    p_jira = doc.add_paragraph(style='List Bullet')
    p_jira.add_run("UC-02 (Risk Score) -> Epic: Risk Engine -> ").bold = True
    p_jira.add_run("Story GLOF-02: Implement hybrid Costa/ML logic in core/risk_engine.py.")
    p_jira = doc.add_paragraph(style='List Bullet')
    p_jira.add_run("UC-03 (Alerts) -> Epic: Notifications -> ").bold = True
    p_jira.add_run("Story GLOF-03: Implement SSEContext.js in Frontend; Set up Celery worker in Backend.")

    # 6. Plan for Next Week
    doc.add_heading("6. Plan for Next Week (Week 3)", level=1)
    doc.add_paragraph("Entering Week 3, the project will move into the Object-Oriented Analysis phase. The primary objectives include:", style='List Bullet')
    doc.add_paragraph("Focus on Static Analysis to identify domain entities and their structural relationships.", style='List Bullet')
    doc.add_paragraph("Develop an Object Diagram to model specific instances and data snapshots.", style='List Bullet')
    doc.add_paragraph("Create the Initial Class Diagram (Domain Model) focusing on attributes and cardinality relationships, deferring method definitions.", style='List Bullet')

    # Footer
    p_footer = doc.add_paragraph("\n\nMentor Signature: __________________      Project Guide Signature: __________________")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(r"c:\GLOF\reports\in_depth", exist_ok=True)
    filename = r"c:\GLOF\reports\in_depth\GLOF_Week_2_Report_InDepth.docx"
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    create_week2_report()
