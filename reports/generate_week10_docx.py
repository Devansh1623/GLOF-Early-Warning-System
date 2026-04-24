import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_week10_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('GLOF Early Warning System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Minor Project Report - Week 10')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Phase 4: Implementation')
    doc.add_paragraph('Focus: Sprint 2 - Integration').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: 06 - 11 April 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_paragraph(doc, 'Week 10 focused entirely on realizing Sprint 2 objectives, which aimed to converge the disjoint backend core system and the visual frontend architecture. We successfully developed and hooked the entire React-based User Interface (UI) to our MongoDB database via Flask endpoints, achieving the final interactive loop. Simultaneously, project tracking through Jira was advanced into agile testing cycles.')
    
    # 2. UI Development and Database Connectivity
    add_heading(doc, '2. UI Development and Database Connectivity', 1)
    add_paragraph(doc, 'Significant headway was made connecting the components built in earlier stages to the central database.')
    
    add_paragraph(doc, 'User Interface (UI) Realization:', bold=True)
    p_ui = doc.add_paragraph(style='List Bullet')
    p_ui.add_run('Dashboard Components: ').bold = True
    p_ui.add_run('Built React pages such as DashboardHome, ChartsPage, and MapPage configured to fetch telemetric statistics directly from the backend.')
    p_ui = doc.add_paragraph(style='List Bullet')
    p_ui.add_run('Advanced Routing: ').bold = True
    p_ui.add_run('The interface now handles seamless failovers with multiple APIs configured depending on the deployment state via utils/api.js logic.')
    
    add_paragraph(doc, '\nDatabase Connections via Endpoints:', bold=True)
    p_db = doc.add_paragraph(style='List Bullet')
    p_db.add_run('Server Sent Events (SSE): ').bold = True
    p_db.add_run('Established resilient asynchronous connections between the UI and backend (utilizing SSEContext.js) forcing UI to respond instantly when environmental sensors exceed Risk Engine thresholds logged within MongoDB.')
    p_db = doc.add_paragraph(style='List Bullet')
    p_db.add_run('Context Providers: ').bold = True
    p_db.add_run('Integrated AuthContext and language I18nContext fully mapped to our persistent schemas for users and preferences.')

    # 3. Jira Agile Workflows
    add_heading(doc, '3. Agile Project Management (Jira Updates)', 1)
    add_paragraph(doc, 'This week required strict tracking synchronization. The Sprint 2 backlog items have been updated to reflect accurate production states:')
    
    p_jira1 = doc.add_paragraph(style='List Bullet')
    p_jira1.add_run('Moved to "In Progress": ').bold = True
    p_jira1.add_run('Tasks related to optimizing specific analytical chart rendering and final error boundary definitions on pages like AdminPage.')
    p_jira1 = doc.add_paragraph(style='List Bullet')
    p_jira1.add_run('Moved to "Testing": ').bold = True
    p_jira1.add_run('Completed feature sets (e.g., AuthPage flow, database read patterns for telemetry) were moved into QA testing phases for intensive bug-hunting by the team.')

    # 4. Plan for Next Phase
    add_heading(doc, '4. Plan for Next Phase (Phase 5: Weeks 11-12)', 1)
    doc.add_paragraph('As Sprint 2 wraps up, we officially step into Phase 5: Testing & Deployment next week.', style='List Bullet')
    doc.add_paragraph('We will conduct extensive End-To-End Testing (E2E).', style='List Bullet')
    doc.add_paragraph('We will lock down environment configurations to initiate our final staging deployments.', style='List Bullet')

    # Ensure output directory exists before saving
    output_dir = r"c:\GLOF\reports\in_depth"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "GLOF_Week_10_Report_InDepth.docx")
    
    try:
        doc.save(file_path)
        print(f"Successfully generated {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_week10_report()
