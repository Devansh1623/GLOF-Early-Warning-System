import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_mermaid_image(mermaid_code, output_path):
    print(f"Fetching diagram for {output_path}...")
    encoded_code = mermaid_code.encode('utf-8')
    import base64
    b64_code = base64.b64encode(encoded_code).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64_code}"
    
    response = requests.get(url)
    if response.status_code == 200:
        with open(output_path, 'wb') as f:
            f.write(response.content)
        return True
    else:
        print(f"Error generating Mermaid image: {response.status_code}")
        print(response.text)
        return False

def generate_week7_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('GLOF Early Warning System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Minor Project Report - Week 7')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Phase 4: Implementation & CI/CD Integration Transition').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: 02-07 March 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Introduction
    add_heading(doc, '1. System Architecture Overview', 1)
    add_paragraph(doc, 'For the seventh week of the project, the primary objective was finalizing the high-level macroscopic view of the System Architecture before commencing active sprint implementations. The architecture dictates how independent functional blocks (Components) communicate with each other and exactly how they will be provisioned onto physical or cloud infrastructure (Nodes).')
    
    add_paragraph(doc, 'The system adopts a decoupled, distributed microservices-inspired architecture. This separation ensures that computationally heavy tasks (like the XGBoost analytical models in the Risk Engine) do not degrade the performance of the outward-facing REST HTTP endpoints or the Server-Sent Event (SSE) connections.')

    # 2. Component Architecture
    add_heading(doc, '2. Component Diagram', 1)
    add_paragraph(doc, 'The physical modules of the GLOF Early Warning System are organized around three primary actors: The IoT Edge Devices, the Cloud Backend, and the Client UI. The following UML Component Diagram visualizes the rigid boundaries between the API, the distinct Database solutions utilized (MongoDB and Redis), and the React UI frontend.')
    
    component_diagram_mermaid = """
    flowchart LR
        subgraph IoT_Edge ["IoT Sensor Arrays"]
            sensor1[Lhonak Lake Sensor]
            sensor2[Teesta River Sensor]
        end

        subgraph Cloud_Infrastructure ["Central Cloud Platform"]
            api[Flask REST API]
            risk[ML Risk Engine Component]
            worker[Celery Background Workers]
            
            db[(MongoDB Persistent Storage)]
            cache[(Redis Message Broker)]
            
            api <--> risk
            api --> db
            api --> cache
            cache <--> worker
        end

        subgraph Client_Tier ["Observer Clients"]
            ui[React Single Page Application]
            email[SMTP Notification Node]
        end

        sensor1 & sensor2 == "HTTP POST (JSON)" ==> api
        cache == "Server-Sent Events" ==> ui
        ui == "HTTP GET/POST" ==> api
        worker == "Resend API" ==> email
    """
    
    comp_img_path = 'c:\\GLOF\\reports\\in_depth\\component_diagram.png'
    os.makedirs('c:\\GLOF\\reports\\in_depth', exist_ok=True)
    if generate_mermaid_image(component_diagram_mermaid, comp_img_path):
        doc.add_picture(comp_img_path, width=Inches(6.0))
        p = doc.add_paragraph('Figure 1: UML Component Diagram showcasing physical modules and inter-component API boundaries.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    else:
        add_paragraph(doc, '[Error: Component Diagram could not be generated]', bold=True)

    doc.add_page_break()

    # 3. Deployment Topology
    add_heading(doc, '3. Deployment Diagram', 1)
    add_paragraph(doc, 'To prepare for the CI/CD pipeline integration scheduled for the upcoming weeks, it was imperative to map the specific software artifacts to their target execution nodes. The application utilizes a PaaS (Platform as a Service) deployment strategy deployed on Render.com, backed by fully-managed database tiers.')
    
    deployment_diagram_mermaid = """
    flowchart TD
        subgraph Render_Network ["Render Cloud Network (VPC)"]
            subgraph Web_Service ["Render Web Service Node"]
                flask(Flask Gunicorn Server)
            end
            
            subgraph Worker_Service ["Render Background Worker Node"]
                celery(Celery Worker Daemon)
            end
            
            subgraph Static_Site ["Render Static Site CDN"]
                nginx(Nginx Web Server)
                react(React Virtual DOM Build)
                nginx --> react
            end
        end

        subgraph External_Managed_Services ["Managed DB & APIs"]
            mongo[(MongoDB Atlas Serverless)]
            upstash[(Upstash Redis Cluster)]
            resend{Resend Mail API Node}
        end

        Client([User Browser]) -- "HTTPS/SSL" --> nginx
        Client -- "HTTPS (SSE Support)" --> flask
        
        flask -- "TCP/TLS" --> mongo
        flask -- "TCP/TLS" --> upstash
        celery -- "TCP/TLS" --> upstash
        celery -- "HTTPS" --> resend
    """
    
    dep_img_path = 'c:\\GLOF\\reports\\in_depth\\deployment_diagram.png'
    if generate_mermaid_image(deployment_diagram_mermaid, dep_img_path):
        doc.add_picture(dep_img_path, width=Inches(6.0))
        p = doc.add_paragraph('Figure 2: UML Deployment Diagram detailing the mapping of system subsystems to Render.com cloud infrastructure and managed database protocols.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    else:
        add_paragraph(doc, '[Error: Deployment Diagram could not be generated]', bold=True)

    # 4. Conclusion
    add_heading(doc, '4. Weekly Conclusion', 1)
    add_paragraph(doc, 'Week 7 successfully finalized the "How" and "Where" of the GLOF architecture. By articulating the logical Components alongside the specific physical hardware Deployment targets (Flask on Render Web Services, React on Render Static Sites, and MongoDB on Atlas VPCs), the development environment is fully documented. The system architecting phase is officially complete.')

    # 5. Plan for Next Week
    add_heading(doc, '5. Plan for Next Week (Week 8)', 1)
    doc.add_paragraph('Entering Week 8, the project will officially commence the execution phases within Sprint 1. The primary objectives include:', style='List Bullet')
    doc.add_paragraph('Develop core application logic (Risk engine models, Flask endpoints) derived strictly from prior sequence and class diagrams.', style='List Bullet')
    doc.add_paragraph('Establish rigorous Git Strategies utilizing feature branching prior to merging into Main.', style='List Bullet')

    try:
        doc.save('c:\\GLOF\\reports\\in_depth\\GLOF_Week_7_Report_InDepth.docx')
        print("Successfully generated c:\\GLOF\\reports\\in_depth\\GLOF_Week_7_Report_InDepth.docx")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_week7_report()
