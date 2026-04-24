import os
import base64
import json
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_mermaid_image(mermaid_code):
    state = {
        "code": mermaid_code,
        "mermaid": {
            "theme": "default"
        }
    }
    b64 = base64.urlsafe_b64encode(json.dumps(state).encode('utf-8')).decode('ascii')
    url = f"https://mermaid.ink/img/{b64}?type=png"
    
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to fetch image: {e}")
        return None

def create_week4_report():
    doc = Document()
    
    # Title
    title = doc.add_heading("Week 4 Project Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run("Project Title: ").bold = True
    p.add_run("GLOF Early Warning System (GLOFWatch)\n")
    p.add_run("Phase 2: ").bold = True
    p.add_run("Object-Oriented Analysis\n")
    p.add_run("Focus: ").bold = True
    p.add_run("Dynamic Analysis: Modeling how the system reacts to events in real-time.\n")
    p.add_run("Date Range: ").bold = True
    p.add_run("09 - 14 February 2026\n")
    
    # 1. Executive Summary & Dynamic Analysis
    doc.add_heading("1. Executive Summary & Dynamic Analysis", level=1)
    doc.add_paragraph(
        "During Week 4, the engineering focus moved to Dynamic Analysis. The objective was to formally map "
        "the complex, real-time business logic required for the GLOF Early Warning System. "
        "Because our architecture utilizes asynchronous tasks and event-driven data streaming, a traditional "
        "linear sequence was insufficient. We utilized UML Activity Diagrams with Fork and Join nodes to "
        "model parallel execution states within the backend."
    )
    
    # 2. System Reaction to Events
    doc.add_heading("2. System Reaction to Disastrous Events", level=1)
    doc.add_paragraph(
        "The most complex workflow in our application occurs when a High or Critical anomaly is detected by the Risk Engine. "
        "To achieve millisecond-level responsiveness, the system executes multiple parallel processes:"
    )
    p_steps = doc.add_paragraph(style='List Bullet')
    p_steps.add_run("Persistence: ").bold = True
    p_steps.add_run("The raw telemetry reading is saved to the MongoDB collection immediately.")
    p_steps = doc.add_paragraph(style='List Bullet')
    p_steps.add_run("In-Memory Pub/Sub: ").bold = True
    p_steps.add_run("The generated Alert object is published to a Redis channel where a Flask listener streams it to all active React Web UI clients via Server-Sent Events (SSE).")
    p_steps = doc.add_paragraph(style='List Bullet')
    p_steps.add_run("Asynchronous Workers: ").bold = True
    p_steps.add_run("A task is pushed to the Celery message broker (backed by Redis) which dispatches an emergency email notification via the Resend API without blocking the main thread.")
    
    # 3. UML Deliverable: Activity Diagram
    doc.add_heading("3. UML Deliverable: Activity Diagram (Telemetry Workflow)", level=1)
    doc.add_paragraph(
        "The Following UML Activity Diagram illustrates the exact parallel execution pathways taken by the backend when an incoming telemetry POST request is received. "
        "Special attention was given to the 'Fork' execution streams when an anomoly is detected."
    )
    
    activity_diagram_code = """
flowchart TD
    Start((Start)) --> Ingest[Receive POST /api/telemetry]
    Ingest --> Validate{Validate Marshmallow Schema}
    
    Validate -- Invalid --> Error[Return 400 Bad Request] --> End((End))
    Validate -- Valid --> Fork1((Fork Execution))
    
    Fork1 --> SaveMongo[Insert Telemetry to MongoDB]
    Fork1 --> CalcRisk[Call risk_engine.calculate_risk]
    
    CalcRisk --> EvalRisk{Assess Risk Tier}
    
    EvalRisk -- Low / Moderate --> Join1((Join Threads))
    EvalRisk -- High / Critical --> TriggerAlert[Generate Alert Object]
    
    TriggerAlert --> Fork2((Fork Alert Dispatch))
    
    Fork2 --> RedisPub[PUBLISH to Redis Channel]
    Fork2 --> CeleryTask[Queue Celery Email Task]
    Fork2 --> SaveMongoAlert[Insert Alert to MongoDB]
    
    RedisPub --> PushSSE[Dispatch SSE to React Dashboard]
    CeleryTask --> SendEmail[Transmit via Resend Mail API]
    
    SaveMongo --> Join1
    PushSSE --> Join1
    SendEmail --> Join1
    SaveMongoAlert --> Join1
    
    Join1 --> Return200[Return 200 OK Response]
    Return200 --> End
"""
    print("Fetching Activity Diagram...")
    img = get_mermaid_image(activity_diagram_code)
    if img:
        doc.add_picture(img, width=Inches(6.0))
    else:
        doc.add_paragraph("[Failed to load Activity Diagram]", style="Intense Quote")

    # 4. Plan for Next Week
    doc.add_heading("4. Plan for Next Week (Week 5)", level=1)
    doc.add_paragraph("Entering Week 5, the project will advance into Phase 3 (Object-Oriented Design). The primary objectives include:", style='List Bullet')
    doc.add_paragraph("Shift focus to the 'How' of the system—establishing architectures, patterns, and component interactions.", style='List Bullet')
    doc.add_paragraph("Perform Interaction Modeling to define object-to-object coordination over time for fulfilling complex use cases.", style='List Bullet')
    doc.add_paragraph("Design a detailed Sequence Diagram mapping the exact lifecycle of an incoming telemetry request.", style='List Bullet')

    # Footer
    p_footer = doc.add_paragraph("\n\nMentor Signature: __________________      Project Guide Signature: __________________")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(r"c:\GLOF\reports\in_depth", exist_ok=True)
    filename = r"c:\GLOF\reports\in_depth\GLOF_Week_4_Report_InDepth.docx"
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    create_week4_report()
