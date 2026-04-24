import os
import base64
import json
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_mermaid_image(mermaid_code):
    """Fetches a PNG from mermaid.ink given a string of Mermaid syntax."""
    # Mermaid ink accepts base64 encoded JSON
    state = {
        "code": mermaid_code,
        "mermaid": {
            "theme": "default"
        }
    }
    b64 = base64.urlsafe_b64encode(json.dumps(state).encode('utf-8')).decode('ascii')
    url = f"https://mermaid.ink/img/{b64}?type=png"
    
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to fetch image: {e}")
        return None

def create_report(week_num, dates, sprint_goal, status, summary, uml_desc, images, git_text, testing_text, goals):
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Week {week_num} Project Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run(f"Project Title: GLOF Early Warning System (GLOFWatch)\n").bold = True
    p.add_run(f"Institute: Shri Vaishnav Institute of Information Technology\n").bold = True
    p.add_run(f"Date Range: {dates}\n").bold = True
    
    # Sections
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(f"Sprint Goal: {sprint_goal}")
    doc.add_paragraph(f"Status: {status}")
    doc.add_paragraph(summary)
    
    doc.add_heading("2. UML & OOAD Progress", level=1)
    doc.add_paragraph(uml_desc)
    
    for img_name, code in images:
        doc.add_heading(img_name, level=2)
        print(f"Generating image for {img_name}...")
        img_buffer = get_mermaid_image(code)
        if img_buffer:
            doc.add_picture(img_buffer, width=Inches(6.0))
        else:
            doc.add_paragraph(f"[Failed to load Mermaid Diagram: {img_name}]")
    
    doc.add_heading("3. Technical Execution", level=1)
    for p_text in git_text:
        doc.add_paragraph(p_text, style="List Bullet")
        
    doc.add_heading("4. Testing Summary", level=1)
    for k, v in testing_text.items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
        
    doc.add_heading("5. Goals for Next Week", level=1)
    for g in goals:
        doc.add_paragraph(g, style="List Bullet")
        
    p = doc.add_paragraph("\n\nMentor Signature: __________________      Project Guide Signature: __________________")
    
    filename = f"c:\\GLOF\\reports\\GLOF_Week_{week_num}_Report.docx"
    doc.save(filename)
    print(f"Saved: {filename}")
    
reports = [
    {
        "week_num": 1,
        "dates": "19-24 January 2026",
        "sprint_goal": "Define problem statement, set up infrastructure, create initial backlog",
        "status": "On Track",
        "summary": "Initiated GLOF Early Warning System. Analyzed CWC/NDMA reports to identify 12 high-risk lakes. Selected stack: Flask, React, MongoDB, Redis, XGBoost, Docker.",
        "uml_desc": "Identified domain entities: 12 Lakes, 7 Historical Events, 5 Stakeholder actors.",
        "images": [],
        "git_text": ["Repo initialized", ".env configured for 12 variables (JWT, Redis, MongoDB)"],
        "testing_text": {"Unit Tests": 0, "Pass Rate": "N/A"},
        "goals": ["Identify actors/use cases", "Create Use Case Diagram"]
    },
    {
        "week_num": 2,
        "dates": "26-31 January 2026",
        "sprint_goal": "Identify actors, use cases, and system boundaries",
        "status": "On Track",
        "summary": "Mapped 13 use cases across telemetry, risk, auth, and alerts. Five main actors identified.",
        "uml_desc": "Completed Use Case Diagram establishing the primary boundaries.",
        "images": [
            ("Use Case Diagram", '''graph TB
subgraph System
UC1[Ingest Telemetry]
UC2[Calculate Risk]
UC3[Fire Alert]
UC4[Register/Login]
end
Sensor --> UC1
UC1 --> UC2
UC2 --> UC3
Admin --> UC4
Admin --> UC3''')
        ],
        "git_text": ["Backlog created with 14 stories", "Sorted by risk (Telemetry P0)"],
        "testing_text": {"Unit Tests": 0},
        "goals": ["Create Object Diagram", "Create Domain Model"]
    },
    {
        "week_num": 3,
        "dates": "02-07 February 2026",
        "sprint_goal": "Identify domain entities and create Domain Model",
        "status": "On Track",
        "summary": "Completed static analysis identifying 8 key entities. Seeded data for 12 CWC lakes and 7 historical events.",
        "uml_desc": "Created Object Diagram and Initial Class Diagram (Attributes only).",
        "images": [
            ("Domain Model", '''classDiagram
class Lake {
    +String id
    +String name
    +Float risk_score
}
class TelemetryReading {
    +Float temperature
    +Float rainfall
    +Float water_level_rise
}
Lake "1" --> "*" TelemetryReading''')
        ],
        "git_text": ["Domain data seeded: core/lake_data.py"],
        "testing_text": {"Unit Tests": 0},
        "goals": ["Model dynamic behavior with Activity Diagram"]
    },
    {
        "week_num": 4,
        "dates": "09-14 February 2026",
        "sprint_goal": "Model dynamic behavior via Activity Diagram",
        "status": "On Track",
        "summary": "Modeled real-time parallel pipeline for ingestion, risk calculation, pub/sub, and async email via celery.",
        "uml_desc": "Activity Diagram created for Telemetry Ingestion.",
        "images": [
            ("Activity Diagram", '''flowchart TD
Start[Sensor POST /api/telemetry] --> Valid{Validate schema}
Valid --> Fork1((Fork))
Fork1 --> CalcRisk[Calculate Risk Score]
CalcRisk --> Classify[Low/Moderate/High/Critical]
Classify --> Pub[Publish via Redis PUB/SUB]
Pub --> End[End]''')
        ],
        "git_text": ["Validated architecture against app.py flow"],
        "testing_text": {"Unit Tests": 0},
        "goals": ["Sequence Diagrams for telemetry and auth"]
    },
    {
        "week_num": 5,
        "dates": "16-21 February 2026",
        "sprint_goal": "Model object interactions for primary use cases",
        "status": "On Track",
        "summary": "Sequence diagrams constructed detailing interactions between Flask routes, MongoDB, Redis, and Celery workers.",
        "uml_desc": "Sequence diagrams for Telemetry, Authentication, and Alert Lifecycle completed.",
        "images": [
            ("Telemetry Sequence", '''sequenceDiagram
Sensor->>Flask: POST /api/telemetry
Flask->>Redis: RPUSH wl_window
Flask->>RiskEngine: calculate_risk()
RiskEngine-->>Flask: score
Flask->>MongoDB: insert telemetry
Flask->>Redis: PUBLISH stream
Flask-->>Sensor: 200 OK''')
        ],
        "git_text": ["Verified state machine logic (OPEN -> ACK -> RESOLVED)"],
        "testing_text": {"Unit Tests": 0},
        "goals": ["Apply SOLID principles to class design"]
    },
    {
        "week_num": 6,
        "dates": "23-28 February 2026",
        "sprint_goal": "Apply SOLID principles and create Detailed Class Diagram",
        "status": "On Track",
        "summary": "Separated concerns into risk calculation, validation, logging, and notifications.",
        "uml_desc": "Detailed Class Diagram outlining methods and a Package Diagram comprising 6 subsystems.",
        "images": [
            ("Package Diagram", '''graph TB
subgraph backend
    core[core/]
    routes[routes/]
    models[models/]
end
subgraph frontend
    pages[pages/]
    utils[utils/]
end
core --> routes
utils --> pages''')
        ],
        "git_text": ["No circular dependecies identified across 7 core modules"],
        "testing_text": {"Unit Tests": 0},
        "goals": ["Component and Deployment Diagrams"]
    },
    {
        "week_num": 7,
        "dates": "02-07 March 2026",
        "sprint_goal": "Visualize physical architecture",
        "status": "On Track",
        "summary": "Mapped software components to Docker execution nodes and Render cloud deployment environment.",
        "uml_desc": "Component Diagram and Deployment Diagram completed.",
        "images": [
            ("Deployment Diagram", '''graph TB
subgraph Cloud Render
    Gunicorn[Web Backend Port 5000]
    Nginx[React Static Site Port 80]
end
subgraph Data
    Atlas[MongoDB]
    Redis[Redis Server]
end
Gunicorn --> Atlas
Gunicorn --> Redis''')
        ],
        "git_text": ["docker-compose.yml config updated", "render.yaml validated"],
        "testing_text": {"Unit Tests": 0},
        "goals": ["Implementation Sprint 1: Core Backend"]
    },
    {
        "week_num": 8,
        "dates": "23-28 March 2026",
        "sprint_goal": "Develop core backend logic",
        "status": "On Track",
        "summary": "Implemented 1,808 lines of Python across 10 modules. Included Risk Engine, Auth Routes, Celery tasks, and JWT.",
        "uml_desc": "Class diagrams updated to reflect actual final methods.",
        "images": [],
        "git_text": ["Branch strategy used: main", "15 backend files created"],
        "testing_text": {"Unit Tests": "15+ completed", "Pass Rate": "70%+"},
        "goals": ["Set up CI/CD pipeline in GitHub Actions"]
    },
    {
        "week_num": 9,
        "dates": "30 March - 04 April 2026",
        "sprint_goal": "Configure GitHub Actions CI",
        "status": "On Track",
        "summary": "Configured 3 jobs: Backend Lint & Test, Frontend Build, and Docker Compose Build.",
        "uml_desc": "CI/CD Pipeline Flow diagram created.",
        "images": [
            ("CI/CD Pipeline Flow", '''flowchart LR
GH[GitHub push] --> J1[Job 1: Backend Pytest]
GH --> J2[Job 2: Frontend Build]
J1 --> J3[Job 3: Docker Compose]
J2 --> J3
J3 --> Deploy[Render.com Deploy]''')
        ],
        "git_text": [".github/workflows/ci.yml added", "Makefile commands registered for automation"],
        "testing_text": {"Unit Tests": "15+", "Pass Rate": "100% locally"},
        "goals": ["Frontend integration for Sprint 2"]
    },
    {
        "week_num": 10,
        "dates": "06-11 April 2026",
        "sprint_goal": "Full-stack integration",
        "status": "On Track",
        "summary": "Connected React frontend (11 pages) with Flask SSE backend. Included AuthContext and SSEContext. Total project size ~7,321 lines.",
        "uml_desc": "Component architecture diagram finalized with Glacial Observatory UI components.",
        "images": [],
        "git_text": ["43 technical files pushed", "Render.com automated deployment functioning"],
        "testing_text": {"Integration Tests": "Manual E2E verification passed"},
        "goals": ["Final presentation preparation"]
    }
]

if __name__ == "__main__":
    for rep in reports:
        create_report(
            rep["week_num"], rep["dates"], rep["sprint_goal"], rep["status"],
            rep["summary"], rep["uml_desc"], rep["images"], rep["git_text"],
            rep["testing_text"], rep["goals"]
        )
