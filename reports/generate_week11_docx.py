import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_week11_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('GLOF Early Warning System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Minor Project Report - Week 11')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Phase 5: Testing & Deployment')
    doc.add_paragraph('Focus: Quality assurance and final delivery.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: 13 - 18 April 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_paragraph(doc, 'Week 11 shifted our focus strictly towards rigorous Quality Assurance (QA). As the core implementations for both the backend application programming interfaces (APIs) and the frontend dashboard were complete, our engineering efforts aimed to battle-test these endpoints. We finalized comprehensive unit testing configurations using PyTest to validate calculation edges, and validated system integrations to eliminate regressions. Crucially, test validations were codified directly into the project Continuous Integration pipeline.')
    
    # 2. Testing Methodologies
    add_heading(doc, '2. Testing Methodologies Executed', 1)
    add_paragraph(doc, 'Several distinct tiers of testing were executed to provide an umbrella of coverage across the technology stack:')
    
    add_paragraph(doc, 'Unit Testing (PyTest):', bold=True)
    p_unit = doc.add_paragraph(style='List Bullet')
    p_unit.add_run('Logic Validation: ').bold = True
    p_unit.add_run('Implemented strict PyTest suites (e.g., test_risk_engine.py). This suite independently parses the mathematical equations validating the Costa 1988 methodologies and boundary limits without needing an active database connection.')
    p_unit = doc.add_paragraph(style='List Bullet')
    p_unit.add_run('Schema Validation: ').bold = True
    p_unit.add_run('Mock objects and payloads were passed to our Marshmallow models (test_schemas.py) protecting routes against malformed upstream JSON objects from the frontend or simulator.')

    add_paragraph(doc, '\nIntegration Testing:', bold=True)
    p_int = doc.add_paragraph(style='List Bullet')
    p_int.add_run('Route Handlers: ').bold = True
    p_int.add_run('Verified the comprehensive flows spanning from network requests directly to the Mock Database and Redis Pub-Sub brokers (test_routes.py).')
    p_int = doc.add_paragraph(style='List Bullet')
    p_int.add_run('Auth Flows: ').bold = True
    p_int.add_run('Ensured Role-Based Access Controls (RBAC) gracefully deny or accept simulated Administrator queries depending on valid JSON Web Tokens (JWT).')

    # 3. CI/CD DevOps Tooling Integration
    add_heading(doc, '3. Coverage Tooling in CI/CD', 1)
    add_paragraph(doc, 'To ensure technical debt remains low over time, the test structures were mapped into our .github/workflows/ci.yml automations.')
    
    p_cicd = doc.add_paragraph(style='List Bullet')
    p_cicd.add_run('Coverage Enforcement: ').bold = True
    p_cicd.add_run('The CI pipeline now invokes `pytest --cov=core --cov=routes --cov-report=term-missing --cov-fail-under=70`. This instructs the server to fail any prospective builds if the code coverage density for our core analytics falls below 70%.')
    p_cicd = doc.add_paragraph(style='List Bullet')
    p_cicd.add_run('Build Protection: ').bold = True
    p_cicd.add_run('Pull Requests dynamically visualize the PyTest output directly within terminal logs, ensuring transparency between developers during code reviews.')

    # 4. Plan for Final Phase
    add_heading(doc, '4. Plan for Project Completion (Week 12)', 1)
    doc.add_paragraph('Next week (Week 12) concludes the project tracking lifecycle as we move toward "Deployment & Final Review". Key targets:', style='List Bullet')
    doc.add_paragraph('Containerization optimizations and orchestrating deployments using Docker configurations.', style='List Bullet')
    doc.add_paragraph('Finalized reporting formats encompassing end-user and administrative documentation.', style='List Bullet')

    # Ensure output directory exists before saving
    output_dir = r"c:\GLOF\reports\in_depth"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "GLOF_Week_11_Report_InDepth.docx")
    
    try:
        doc.save(file_path)
        print(f"Successfully generated {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_week11_report()
