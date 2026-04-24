import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_mermaid_image(mermaid_code, output_path):
    print(f"Fetching diagram for {output_path}...")
    encoded_code = mermaid_code.encode('utf-8')
    import base64
    b64_code = base64.b64encode(encoded_code).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64_code}"
    
    response = requests.get(url)
    if response.status_code == 200:
        with open(output_path, 'wb') as f:
            f.write(response.content)
        return True
    else:
        print(f"Error generating Mermaid image: {response.status_code}")
        print(response.text)
        return False

def generate_week6_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('GLOF Early Warning System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Minor Project Report - Week 6')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Date: 23-28 February 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Introduction
    add_heading(doc, '1. Detailed Class Design & Packaging', 1)
    add_paragraph(doc, 'For the sixth week, the project focus shifted entirely to software architecture and structural design. The goal was to refine the high-level domain models established in Week 3 into detailed, implementation-ready class structures. This phase involved strictly applying SOLID principles and established design patterns to ensure the GLOF Early Warning System remains scalable, maintainable, and robust against future changes.')
    
    add_paragraph(doc, 'This report presents the system\'s internal structure through detailed Class Diagrams encompassing full method signatures and visibility scoping. Furthermore, a Package Diagram is provided to illustrate the logical separation of concerns into distinct micro-modules and subsystems, facilitating an organized codebase and clear dependency management.')

    # 2. Application of SOLID Principles & Design Patterns
    add_heading(doc, '2. SOLID Principles & Design Patterns Implementation', 1)
    
    add_paragraph(doc, 'Single Responsibility Principle (SRP):', bold=True)
    add_paragraph(doc, 'Strict separation of concerns has been enforced. For example, risk calculation mathematics (contained in the RiskEngine static class) operates entirely independently from MongoDB persistence layers or the Flask routing controllers.')
    
    add_paragraph(doc, 'Dependency Inversion Principle (DIP):', bold=True)
    add_paragraph(doc, 'High-level modules, such as the prediction pipelines, do not depend directly on low-level database operations. Instead, data abstraction layers and Data Access Objects (DAOs) are utilized.')
    
    add_paragraph(doc, 'Design Patterns Applied:', bold=True)
    add_paragraph(doc, '• Singleton Pattern: Utilized for establishing unified connection pools to MongoDB and Redis to prevent connection overhead during high-frequency telemetry ingestion.\n'
                       '• Observer / Pub-Sub Pattern: Deeply integrated via Redis, where the RiskEngine acts as the publisher of high-risk events, and the SSE controllers act as subscribers seamlessly pushing alerts to connected frontend clients.\n'
                       '• Strategy Pattern: Employed within the risk calculation engine to allow dynamic swapping between different algorithmic assessment strategies (e.g., empirical Costa formulas vs. ML-based Isolation Forest anomaly detection).')

    # 3. Detailed Class Diagram
    add_heading(doc, '3. Detailed Class Diagram', 1)
    add_paragraph(doc, 'The following UML Class Diagram provides an exhaustive view of the core system entities. It explicitly details attributes and operations, including parameter types, return types, and strict visibility scopes (Private [-], Public [+], Protected [#]).')
    
    class_diagram_mermaid = """
    classDiagram
        class TelemetryData {
            -String sensorId
            -Float waterLevel
            -Float temperature
            -Float flowRate
            -DateTime timestamp
            +validateRanges() Boolean
            +sanitizeInputs() TelemetryData
            +serialize() JSON
        }
        
        class RiskEngine {
            <<Singleton>>
            -XGBoostModel mlModel
            -Float baselineThreshold
            +calculateTotalRisk(TelematryData data, LakeConfig config) Float
            -applyCostaFormula(Float depth, Float volume) Float
            -detectAnomalies(TelemetryData data) Boolean
        }
        
        class LakeConfig {
            -ObjectId _id
            -String lakeName
            -Dictionary coordinates
            -Float baseVolume
            -Float maxSafeLevel
            +updateThresholds(Dictionary newThresholds) Boolean
            +getAreaMetadata() Dictionary
        }
        
        class NotificationService {
            -ResendClient emailClient
            -RedisConnection redisPool
            +dispatchUrgentAlert(LakeConfig lake, Float riskScore) Void
            -publishSSEEvent(EventPayload payload) Integer
            -queueEmailTask(String recipient, String template) String
        }
        
        class TelemetryAPIController {
            +POST /api/telemetry() Response
            -authenticateSensor(String apiKey) Boolean
            -rateLimitCheck(String ipAddress) Boolean
        }
        
        TelemetryAPIController ..> TelemetryData : "Creates & Validates"
        TelemetryAPIController --> RiskEngine : "Invokes Calculation"
        RiskEngine --> LakeConfig : "Fetches Baseline"
        RiskEngine ..> NotificationService : "Triggers on High Risk"
    """
    
    class_img_path = 'c:\\GLOF\\reports\\in_depth\\detailed_class_diagram.png'
    os.makedirs('c:\\GLOF\\reports\\in_depth', exist_ok=True)
    if generate_mermaid_image(class_diagram_mermaid, class_img_path):
        doc.add_picture(class_img_path, width=Inches(6.0))
        p = doc.add_paragraph('Figure 1: Detailed Class Diagram showing structural logic, relationships, and component visibility.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    else:
        add_paragraph(doc, '[Error: Detailed Class Diagram could not be generated]', bold=True)

    doc.add_page_break()

    # 4. Package/Subsystem Organization
    add_heading(doc, '4. Package Diagram: Logical Subsystems', 1)
    add_paragraph(doc, 'The application architecture diverges from monolithic design, segmenting logic into discrete, cohesive packages. The following UML Package Diagram illustrates how functionality is logically grouped into manageable modules and the dependencies spanning across them.')
    add_paragraph(doc, 'The "backend" is structurally divided into routing layers (Controllers), business logic execution (Core), task dispatching (Workers), and persistence management (Database). The "frontend" aligns strictly with modern React philosophies, bundling components, contextual state providers, and external API services.')

    package_diagram_mermaid = """
    classDiagram
        namespace GLOF_Backend {
            class routes {
                <<Package>>
                auth_router.py
                telemetry_router.py
                alert_router.py
            }
            class core {
                <<Package>>
                risk_engine.py
                ml_inference.py
                schemas.py
            }
            class tasks {
                <<Package>>
                celery_worker.py
                email_dispatcher.py
            }
            class db {
                <<Package>>
                mongo_client.py
                redis_pool.py
                indexes.py
            }
        }
        
        namespace GLOF_Frontend {
            class components {
                <<Package>>
                AlertBanner.jsx
                TelemetryChart.jsx
                MapOverlay.jsx
            }
            class contexts {
                <<Package>>
                AuthContext.jsx
                SSEEventContext.jsx
            }
            class views {
                <<Package>>
                DashboardView.jsx
                SettingsView.jsx
            }
        }
        
        routes --> core : "Delegates Logic"
        core --> db : "Data Access"
        core --> tasks : "Queues Async Jobs"
        views --> components : "Renders"
        views --> contexts : "Consumes State"
        contexts --> routes : "API/SSE Calls"
    """
    
    pkg_img_path = 'c:\\GLOF\\reports\\in_depth\\package_diagram.png'
    if generate_mermaid_image(package_diagram_mermaid, pkg_img_path):
        doc.add_picture(pkg_img_path, width=Inches(6.0))
        p = doc.add_paragraph('Figure 2: Package Diagram illustrating the structural organization of backend modules and frontend component boundaries.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    else:
        add_paragraph(doc, '[Error: Package Diagram could not be generated]', bold=True)


    # 5. Conclusion
    add_heading(doc, '5. Weekly Conclusion', 1)
    add_paragraph(doc, 'Week 6 concluded with a deeply granular definition of the project’s structural components. By formalizing class interfaces and rigorously applying SOLID design principles, the logical blueprint is solidified. Furthermore, explicitly modeling the architectural packaging has clarified subsystem responsibilities. The system is now thoroughly modeled from high-level abstract boundaries down to internal method visibilities, establishing an unshakable concrete foundation to finalize the integration and deployment efforts in the impending weeks.')

    # 6. Plan for Next Week
    add_heading(doc, '6. Plan for Next Week (Week 7)', 1)
    add_paragraph(doc, 'Entering Week 7, the focus will transition to holistic System Architecture before moving into the implementation phase. The primary objectives include:')
    add_paragraph(doc, '• Component Diagram: Visualizing the physical modules of the system (API server, MongoDB cluster, Redis cache, React Web UI) and their high-level interfaces.')
    add_paragraph(doc, '• Deployment Diagram: Mapping the software components to hardware and cloud infrastructure nodes, preparing for CI/CD integration and deployment onto platform-as-a-service providers.')

    try:
        doc.save('c:\\GLOF\\reports\\in_depth\\GLOF_Week_6_Report_InDepth.docx')
        print("Successfully generated c:\\GLOF\\reports\\in_depth\\GLOF_Week_6_Report_InDepth.docx")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_week6_report()
