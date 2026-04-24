import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return
        
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    
    for line in lines:
        line = line.rstrip('\n')
        
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            p = doc.add_paragraph()
            if in_code_block:
                p.add_run("--- Code / Diagram block begins ---").italic = True
            else:
                p.add_run("--- Code / Diagram block ends ---").italic = True
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            # try to simulate code block styling if possible, else just normal
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue
            
        if not line:
            # Empty line
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(line[line.find(' ')+1:].strip(), style='List Number')
        else:
            p = doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Saved {docx_path}")

md_files = [
    (r"C:\Users\devan\.gemini\antigravity\brain\8732c192-8c31-4de4-b0b2-5cc90287e941\artifacts\glof_uml_data_synthesis.md", r"c:\GLOF\reports\GLOF_UML_Data_Synthesis.docx"),
    (r"C:\Users\devan\.gemini\antigravity\brain\8732c192-8c31-4de4-b0b2-5cc90287e941\artifacts\glof_diagram_prompts.md", r"c:\GLOF\reports\GLOF_Diagram_Prompts.docx")
]

for md, docx in md_files:
    convert_md_to_docx(md, docx)
