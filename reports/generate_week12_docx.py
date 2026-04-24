import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_week12_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('GLOF Early Warning System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Final Project Report - Week 12')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Phase 5: Deployment & Final Review')
    doc.add_paragraph('Focus: Staging, Production Handover, and Demonstration.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: 20 - 25 April 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_paragraph(doc, 'Week 12 marks the culmination of the GLOF Early Warning System project. With the application previously rigorously verified through CI/CD pipelines and holistic UI-Database interactions, this final phase focuses solely on bringing the application into a public staging environment. Utilizing Render Cloud Services alongside robust Docker architectures, we finalized the delivery format, created robust handover documentation, and executed final presentations.')
    
    # 2. Production Environments
    add_heading(doc, '2. Production & Staging Environments', 1)
    add_paragraph(doc, 'The application architecture formally graduated from local development servers to remote infrastructure. We established the production footprint using Infrastructure-as-Code via `render.yaml`.')
    
    add_paragraph(doc, 'Cloud Staging Resources:', bold=True)
    p_cloud = doc.add_paragraph(style='List Bullet')
    p_cloud.add_run('Backend API Provisioning (glof-ews-api): ').bold = True
    p_cloud.add_run('Flask environments run via Gunicorn (`start.sh`) configured alongside strict environment variable isolation. Sensitive credentials (MONGO_URI, JWT_SECRET_KEY, RESEND_API_KEY) are decoupled dynamically at deployment.')
    p_cloud = doc.add_paragraph(style='List Bullet')
    p_cloud.add_run('Frontend Edge Provisioning (glof-frontend): ').bold = True
    p_cloud.add_run('The React SPA was compiled statically using `npm run build` and optimized with rewrite rules to ensure seamless client-side routing. It is configured to intelligently fetch against the exact API endpoint provisioned during backend build.')

    # 3. Deployment Design Mapping
    add_heading(doc, '3. Deployment Diagram Finalization', 1)
    add_paragraph(doc, 'Our formal Deployment Diagram, introduced earlier in the project lifecycle, was updated to perfectly reflect the real-world deployment on Render. The static site servers now route to Python Worker nodes via REST, which securely pass analytics over to managed MongoDB Atlas clusters and Upstash Redis event busses.')

    # 4. Handover & Demonstration
    add_heading(doc, '4. Project Demo and Handover Verification', 1)
    add_paragraph(doc, 'The final deliverables were explicitly crafted to guarantee a smooth transition for stakeholders:')
    
    p_ho1 = doc.add_paragraph(style='List Bullet')
    p_ho1.add_run('Docker Containers: ').bold = True
    p_ho1.add_run('We maintained an isolated cross-platform `docker-compose.yml` that flawlessly ties the database and API together to allow evaluators isolated evaluation runs with single commands.')
    p_ho1 = doc.add_paragraph(style='List Bullet')
    p_ho1.add_run('Comprehensive README: ').bold = True
    p_ho1.add_run('The codebase root features a massive, well-maintained `README.md` consisting of sequence mapping, installation instructions, database schemas, and explicit ML configurations.')
    
    add_paragraph(doc, '\nThe GLOF Early Warning System concludes successful execution across all software engineering phases.')

    # Ensure output directory exists before saving
    output_dir = r"c:\GLOF\reports\in_depth"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "GLOF_Week_12_Report_InDepth.docx")
    
    try:
        doc.save(file_path)
        print(f"Successfully generated {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_week12_report()
