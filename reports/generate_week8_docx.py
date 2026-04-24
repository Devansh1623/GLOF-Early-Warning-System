import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_week8_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('GLOF Early Warning System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Minor Project Report - Week 8')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Phase 4: Implementation & CI/CD Integration').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Focus: Coding, version control, and automated pipelines').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: 23 - 28 March 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary: Sprint 1', 1)
    add_paragraph(doc, 'Week 8 marks the official transition from system design into active development, kicking off Sprint 1. The primary focus of this sprint was the implementation of core backend features directly derived from the Class and Sequence diagrams established in previous weeks. Furthermore, strict version control practices were instantiated to manage collaboration effectively.')
    
    # 2. Core Feature Implementation
    add_heading(doc, '2. Core Logic Implementation', 1)
    add_paragraph(doc, 'Following the SOLID principles established in Week 6, the core business logic has been developed into distinct, modular Python packages within the Flask environment.')
    
    add_paragraph(doc, 'Risk Engine (risk_engine.py):', bold=True)
    add_paragraph(doc, 'The mathematical backbone of the system was implemented. It features a hybrid calculation model that utilizes XGBoost and Isolation Forest models for machine learning-based anomaly detection, while maintaining a robust fallback to empirical formulas (incorporating temperature, rainfall, and water level changes based on the Costa 1988 approach).')

    add_paragraph(doc, 'Authentication & Security (auth.py):', bold=True)
    add_paragraph(doc, 'To secure the REST APIs, Role-Based Access Control (RBAC) was implemented using JSON Web Tokens (JWT). Registration and login controllers were built alongside a password reset flow integrated with the Resend Email API. Strict JSON schema validation using Marshmallow ensures data integrity before database interaction.')
    
    add_paragraph(doc, 'Telemetry Simulation (mock_telemetry.py):', bold=True)
    add_paragraph(doc, 'To test the pipeline without physical sensors, a simulator was developed. This simulator fetches real-time meteorological data via the Open-Meteo API and calculates diurnal temperature drifts and randomized water level spikes to mimic genuine environmental behaviors.')

    # 3. Version Control Strategy
    add_heading(doc, '3. Git Branching Strategy', 1)
    add_paragraph(doc, 'To maintain a clean and reliable codebase as complexity increases, a formalized Git workflow has been strictly enforced across the team.')
    
    add_paragraph(doc, 'Feature Branching Workflow:', bold=True)
    add_paragraph(doc, '• Main Branch (\'main\'): Represents the production-ready state. Direct commits to this branch are strictly prohibited.')
    add_paragraph(doc, '• Develop Branch (\'develop\'): Serves as the primary integration branch for the next release.')
    add_paragraph(doc, '• Feature Branches: Individual tasks (e.g., \'feat/risk-engine\', \'feat/auth-jwt\') are branched off \'develop\'. Once the feature is complete and tested, a Pull Request is opened to merge the code back into \'develop\'.')
    
    add_paragraph(doc, 'This hierarchical structure guarantees that the \'main\' branch remains stable for continuous deployment, encapsulating work cleanly based on Agile Jira tickets.')

    # 4. Plan for Next Week
    add_heading(doc, '4. Plan for Next Week (Week 9)', 1)
    doc.add_paragraph('Entering Week 9, the project will advance into Sprint 2. The primary objectives include:', style='List Bullet')
    doc.add_paragraph('Integration of Data Sources: Connecting the frontend React UI hooks deeply with the established Flask REST endpoints and Server-Sent Events (SSE) streams.', style='List Bullet')
    doc.add_paragraph('Development of Automated Pipelines: Setting up GitHub Actions or similar automated triggers for Continuous Integration testing.', style='List Bullet')

    # Ensure output directory exists before saving
    output_dir = r"c:\GLOF\reports\in_depth"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "GLOF_Week_8_Report_InDepth.docx")
    
    try:
        doc.save(file_path)
        print(f"Successfully generated {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_week8_report()
