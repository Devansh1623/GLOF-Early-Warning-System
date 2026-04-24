import os
import base64
import json
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_mermaid_image(mermaid_code):
    state = {
        "code": mermaid_code,
        "mermaid": {
            "theme": "default"
        }
    }
    b64 = base64.urlsafe_b64encode(json.dumps(state).encode('utf-8')).decode('ascii')
    url = f"https://mermaid.ink/img/{b64}?type=png"
    
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to fetch image: {e}")
        return None

def create_week3_report():
    doc = Document()
    
    # Title
    title = doc.add_heading("Week 3 Project Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run("Project Title: ").bold = True
    p.add_run("GLOF Early Warning System (GLOFWatch)\n")
    p.add_run("Phase 2: ").bold = True
    p.add_run("Object-Oriented Analysis\n")
    p.add_run("Focus: ").bold = True
    p.add_run("Static Analysis: Domain entities and their relationships.\n")
    p.add_run("Date Range: ").bold = True
    p.add_run("02 - 07 February 2026\n")
    
    # 1. Executive Summary & Static Analysis
    doc.add_heading("1. Executive Summary & Static Analysis", level=1)
    doc.add_paragraph(
        "During Week 3, the focus shifted to understanding the 'What' of the system. We performed a Static Analysis "
        "to identify the core domain entities based strictly on the data requirements of our real-time architecture, "
        "specifically looking at the MongoDB collections defined in the codebase. The relationships between these entities "
        "were mapped without dwelling on operational methods."
    )
    
    # 2. Identified Domain Entities
    doc.add_heading("2. Identifying Domain Entities", level=1)
    doc.add_paragraph(
        "Based on the schemas established in `backend/core/schemas.py` and our database structure, four primary entities were isolated:"
    )
    
    p_entities = doc.add_paragraph(style='List Bullet')
    p_entities.add_run("Lake: ").bold = True
    p_entities.add_run("Represents a physical glacial lake being monitored (e.g., South Lhonak Lake).")
    p_entities = doc.add_paragraph(style='List Bullet')
    p_entities.add_run("Telemetry: ").bold = True
    p_entities.add_run("Time-series sensor data points including water_level, temperature, and rainfall.")
    p_entities = doc.add_paragraph(style='List Bullet')
    p_entities.add_run("Alert: ").bold = True
    p_entities.add_run("System-generated warning objects triggered when High or Critical risks are reached.")
    p_entities = doc.add_paragraph(style='List Bullet')
    p_entities.add_run("User: ").bold = True
    p_entities.add_run("System administrators defining Lake thresholds and managing subscriber access.")

    # 3. UML Deliverable: Initial Class Diagram (Domain Model)
    doc.add_heading("3. UML Deliverable: Domain Model (Initial Class Diagram)", level=1)
    doc.add_paragraph(
        "This Initial Class Diagram outlines the properties of the entities and their cardinalities (1:1, 1:N). "
        "As per the Phase 2 requirements, methods are excluded to focus purely on the static data relationships."
    )
    
    class_diagram_code = """
classDiagram
    class Lake {
        +String _id
        +String name
        +Float base_depth
        +Float max_depth
        +Float area_km2
        +Float risk_score
        +String status
    }
    class Telemetry {
        +String _id
        +String lake_id
        +DateTime timestamp
        +Float temperature
        +Float rainfall
        +Float water_level
        +Float water_level_rise
    }
    class Alert {
        +String _id
        +String lake_id
        +String risk_tier
        +String message
        +DateTime timestamp
        +String status
    }
    class User {
        +String _id
        +String username
        +String password_hash
        +String role
    }

    Lake "1" --> "*" Telemetry : generates
    Lake "1" --> "*" Alert : triggers
"""
    print("Fetching Domain Model...")
    img1 = get_mermaid_image(class_diagram_code)
    if img1:
        doc.add_picture(img1, width=Inches(6.0))
    else:
        doc.add_paragraph("[Failed to load Domain Model Diagram]", style="Intense Quote")

    # 4. UML Deliverable: Object Diagram
    doc.add_heading("4. UML Deliverable: Object Diagram", level=1)
    doc.add_paragraph(
        "To contextualize the Domain Model, the following Object Diagram details a specific snapshot in time. "
        "It models an instance of 'South Lhonak Lake' linked with a specific critical telemetry reading that has triggered an Alert instance."
    )

    object_diagram_code = """
classDiagram
    class Lake_Instance {
        <<Instance : Lake>>
        _id = "lk_south_lhonak_01"
        name = "South Lhonak Lake"
        base_depth = 45.0
        risk_score = 92.5
        status = "CRITICAL"
    }
    
    class Telemetry_Instance {
        <<Instance : Telemetry>>
        lake_id = "lk_south_lhonak_01"
        timestamp = "2026-02-05T14:30:00Z"
        water_level = 47.5
        water_level_rise = 2.5
        rainfall = 112.4
    }
    
    class Alert_Instance {
        <<Instance : Alert>>
        lake_id = "lk_south_lhonak_01"
        risk_tier = "CRITICAL"
        status = "ACTIVE"
        message = "Imminent GLOF risk detected."
    }

    Lake_Instance "1" --> "1..*" Telemetry_Instance : generated
    Lake_Instance "1" --> "1" Alert_Instance : triggered
"""
    print("Fetching Object Diagram...")
    img2 = get_mermaid_image(object_diagram_code)
    if img2:
        doc.add_picture(img2, width=Inches(6.0))
    else:
        doc.add_paragraph("[Failed to load Object Diagram]", style="Intense Quote")

    # 5. Plan for Next Week
    doc.add_heading("5. Plan for Next Week (Week 4)", level=1)
    doc.add_paragraph("Entering Week 4, the project will continue the Object-Oriented Analysis phase but shift to Dynamic Analysis. The primary objectives include:", style='List Bullet')
    doc.add_paragraph("Model how the system reacts to complex real-time events, particularly high-risk telemetry spikes.", style='List Bullet')
    doc.add_paragraph("Develop an Activity Diagram to visualize the parallel workflows, conditional routing, and asynchronous business logic.", style='List Bullet')

    # Footer
    p_footer = doc.add_paragraph("\n\nMentor Signature: __________________      Project Guide Signature: __________________")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(r"c:\GLOF\reports\in_depth", exist_ok=True)
    filename = r"c:\GLOF\reports\in_depth\GLOF_Week_3_Report_InDepth.docx"
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    create_week3_report()
