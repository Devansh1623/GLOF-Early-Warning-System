import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def update_docx():
    template_path = r"C:\GLOF\reports\report format.docx"
    output_path = r"C:\GLOF\reports\in_depth\GLOF_Week10_Report.docx"
    
    if not os.path.exists(template_path):
        print(f"Error: Template not found at {template_path}")
        return

    doc = Document(template_path)

    # --- Update Header Tables ---
    
    # Update Table 0 (Main Title)
    if len(doc.tables) > 0:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                if 'Week 8' in cell.text:
                    cell.text = cell.text.replace('Week 8: Sprint 1 – Core Feature Development', 'Week 10: Sprint 2 – Integration')
                    cell.text = cell.text.replace('Week 8: Sprint 1 - Core Feature Development', 'Week 10: Sprint 2 - Integration')

    # Update Table 1 (Metadata Table)
    if len(doc.tables) > 1:
        table = doc.tables[1]
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                if key == 'Report Week':
                    row.cells[1].text = 'Week 10 – Sprint 2: Integration'
                elif key == 'Reporting Period':
                    row.cells[1].text = '06 April 2026 – 11 April 2026'
                elif key == 'Phase':
                    row.cells[1].text = 'Phase 4: Implementation & CI/CD Integration'
                elif key == 'Jira Story':
                    row.cells[1].text = 'CPG-40 – Week 10: Sprint 2 – Integration'

    # --- Delete old content ---
    delete_flag = False
    for p in doc.paragraphs:
        if p.text.startswith("1. Objectives") or p.text.startswith("1. Executive"):
            delete_flag = True
        
        if delete_flag:
            # Remove paragraph XML element
            p._element.getparent().remove(p._element)

    # --- Helper Functions ---
    def add_heading(text, level):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        return p

    # --- Inject Week 10 Content ---
    add_heading('1. Executive Summary: Sprint 2 Integration', 1)
    add_paragraph('Week 10 focused heavily on integrating the frontend and backend of the GLOF Early Warning System. The primary objectives were developing the User Interface (UI) components and establishing stable connections to the MongoDB database. We also progressed into the "Testing" phase for several completed Jira stories.')

    add_heading('2. User Interface Development', 1)
    add_paragraph('React was utilized to build dynamic dashboard components, including the Events Timeline page, which provides a tabular view of historical GLOF incidents across various regions with dynamic filtering.')

    add_heading('2.1 Historical Events Page (EventsPage.js)', 2)
    add_paragraph('The frontend now fetches data from the Flask REST endpoints securely, using an authentication wrapper (authFetch) to include JWT tokens. Users can filter events by State and Severity.')

    add_paragraph('Code Highlight: Dynamic Filtering in React', bold=True)
    code_ui = '''useEffect(() => {
    let url = '/api/events/?';
    if (stateFilter) url += `state=${encodeURIComponent(stateFilter)}&`;
    if (sevFilter)   url += `severity=${encodeURIComponent(sevFilter)}&`;
    
    authFetch(url)
      .then(r => r.json())
      .then(data => {
        if (Array.isArray(data)) setEvents(data);
      })
      .catch(() => {});
}, [stateFilter, sevFilter]);'''
    add_code_block(code_ui)

    add_heading('3. Database Connection & Management', 1)
    add_paragraph('The backend system successfully establishes and manages persistent connections to the MongoDB instance. In test environments, a mongomock database is utilized to prevent polluting production data.')

    add_heading('3.1 Flask MongoDB Integration (app.py)', 2)
    add_paragraph('The PyMongo client was set up to handle data ingestion from IoT endpoints, persisting telemetry alongside calculation metrics from the Risk Engine.')

    add_paragraph('Code Highlight: Database and Index Setup', bold=True)
    code_db = '''from pymongo import MongoClient, DESCENDING

MONGO_URI = os.environ.get("MONGO_URI", "mongodb://localhost:27017/glof_db")
mongo_client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=5000)
db = mongo_client.glof_db

# Example Telemetry Persistence
tel_doc = {
    "lake_id": lake_id,
    "lake_name": data["lake_name"],
    "timestamp": timestamp,
    "temperature": temp,
    "water_level_rise": wl_rise,
    "risk_score": risk["score"],
    "risk_level": risk["level"]
}
db.telemetry.insert_one(tel_doc)'''
    add_code_block(code_db)

    add_heading('4. Agile Project Management', 1)
    add_paragraph('As Sprint 2 progressed, project tracking via Jira was rigorously maintained. Several user stories related to API integration and Database setup were successfully moved from the "To Do" to the "In Progress" and "Testing" columns as they entered the QA phase of the pipeline.')

    add_heading('5. Plan for Phase 5 (Weeks 11-12)', 1)
    add_paragraph('The upcoming weeks will signify the transition into Phase 5: Testing & Deployment. The core tasks will include setting up a production deployment utilizing Docker, and verifying end-to-end functionality including automated alerts and notification mechanisms.')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Week 10 report successfully generated at: {output_path}")

if __name__ == "__main__":
    update_docx()
