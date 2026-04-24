import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def update_docx():
    template_path = r"C:\GLOF\reports\report format.docx"
    output_path = r"C:\GLOF\reports\in_depth\GLOF_Week11_Report.docx"
    
    if not os.path.exists(template_path):
        print(f"Error: Template not found at {template_path}")
        return

    doc = Document(template_path)

    # --- Update Header Tables ---
    
    # Update Table 0 (Main Title)
    if len(doc.tables) > 0:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                # Replace Week 8/10 placeholders with Week 11
                if 'Week' in cell.text:
                    # Generic replace for any "Week X" pattern in that specific title block
                    import re
                    cell.text = re.sub(r'Week \d+.*', 'Week 11: Testing Phase', cell.text)

    # Update Table 1 (Metadata Table)
    if len(doc.tables) > 1:
        table = doc.tables[1]
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                if key == 'Report Week':
                    row.cells[1].text = 'Week 11 – Testing Phase'
                elif key == 'Reporting Period':
                    row.cells[1].text = '13 April 2026 – 18 April 2026'
                elif key == 'Phase':
                    row.cells[1].text = 'Phase 5: Testing & Deployment'
                elif key == 'Jira Story':
                    row.cells[1].text = 'CPG-42 – Week 11: Testing Phase'

    # --- Delete old content ---
    delete_flag = False
    for p in doc.paragraphs:
        if p.text.startswith("1. Objectives") or p.text.startswith("1. Executive"):
            delete_flag = True
        
        if delete_flag:
            p._element.getparent().remove(p._element)

    # --- Helper Functions ---
    def add_heading(text, level):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        return p

    # --- Inject Week 11 Content ---
    add_heading('1. Executive Summary: Testing Phase', 1)
    add_paragraph('Week 11 marks the transition into Phase 5: Testing & Deployment. The primary focus of this week was Quality Assurance (QA). We implemented a comprehensive testing suite comprising unit and integration tests to ensure that every system component performs as intended before the final production rollout.')

    add_heading('2. Unit Testing (PyTest)', 1)
    add_paragraph('Unit tests were developed to validate the internal logic of individual modules, specifically the Risk Engine and the data validation schemas. These tests ensure that the mathematical models remain accurate and that edge cases (e.g., extreme weather readings) are handled gracefully.')

    add_paragraph('Code Highlight: Risk Engine Unit Test', bold=True)
    code_unit = '''def test_calculate_risk_critical():
    # Test for critical conditions: high temp, high rain, high rise
    result = calculate_risk(temperature=28.0, rainfall=120.0, water_level_rise=250.0)
    assert result["level"] == "Critical"
    assert result["score"] >= 80

def test_calculate_risk_stable():
    # Test for normal/stable conditions
    result = calculate_risk(temperature=5.0, rainfall=0.0, water_level_rise=10.0)
    assert result["level"] == "Low"
    assert result["score"] < 20'''
    add_code_block(code_unit)

    add_heading('3. Integration Testing', 1)
    add_paragraph('Integration tests were implemented to verify the data flow between components—specifically between the REST API controllers and the MongoDB database. We utilized "mongomock" to simulate database interactions, allowing us to test route behavior without side effects on production data.')

    add_paragraph('Code Highlight: API Route Integration Test', bold=True)
    code_integration = '''def test_get_lakes_authenticated(client, auth_header):
    # Verify that the /api/lakes endpoint returns 200 with valid JWT
    response = client.get("/api/lakes/", headers=auth_header)
    assert response.status_code == 200
    data = response.get_json()
    assert isinstance(data, list)

def test_telemetry_ingestion_valid(client):
    # Verify sensor telemetry ingestion
    payload = {
        "lake_id": "GL001", "lake_name": "Test Lake",
        "temperature": 10.5, "rainfall": 5.0, "water_level_rise": 20.0
    }
    response = client.post("/api/telemetry", json=payload)
    assert response.status_code == 200'''
    add_code_block(code_integration)

    add_heading('4. Tooling & CI/CD Integration', 1)
    add_paragraph('To maintain quality across all future commits, the testing reports were fully integrated into our GitHub Actions CI/CD pipeline. The pipeline is configured to fail the build if the backend test coverage drops below 70%, ensuring that no untested code reaches the main branches.')

    add_heading('5. Plan for Week 12 (Final Review)', 1)
    add_paragraph('As we enter the final week, the focus will shift to production deployment using Docker and the final review of system documentation. We will finalize the deployment diagrams and conduct an end-to-end system walkthrough to verify the GLOF Early Warning System is ready for delivery.')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Week 11 report successfully generated at: {output_path}")

if __name__ == "__main__":
    update_docx()
