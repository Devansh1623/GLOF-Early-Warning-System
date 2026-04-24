import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def update_docx():
    template_path = r"C:\GLOF\reports\report format.docx"
    output_path = r"C:\GLOF\reports\in_depth\GLOF_Week12_Report.docx"
    
    if not os.path.exists(template_path):
        print(f"Error: Template not found at {template_path}")
        return

    doc = Document(template_path)

    # --- Update Header Tables ---
    
    # Update Table 0 (Main Title)
    if len(doc.tables) > 0:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                if 'Week' in cell.text:
                    cell.text = re.sub(r'Week \d+.*', 'Week 12: Deployment & Final Review', cell.text)

    # Update Table 1 (Metadata Table)
    if len(doc.tables) > 1:
        table = doc.tables[1]
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                if key == 'Report Week':
                    row.cells[1].text = 'Week 12 – Final Phase'
                elif key == 'Reporting Period':
                    row.cells[1].text = '20 April 2026 – 25 April 2026'
                elif key == 'Phase':
                    row.cells[1].text = 'Phase 5: Testing & Deployment (Final)'
                elif key == 'Jira Story':
                    row.cells[1].text = 'CPG-50 – Week 12: Production Deployment'

    # --- Delete old content ---
    # The template has placeholder content starting with "1. Objectives" or similar.
    # We remove everything from that point onwards to inject new content.
    delete_flag = False
    for p in doc.paragraphs:
        if p.text.startswith("1. Objectives") or p.text.startswith("1. Executive") or p.text.startswith("1. Introduction"):
            delete_flag = True
        
        if delete_flag:
            p._element.getparent().remove(p._element)

    # --- Helper Functions ---
    def add_heading(text, level):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        return p

    # --- Inject Week 12 Content ---
    add_heading('1. Executive Summary: Production Readiness', 1)
    add_paragraph('Week 12 represents the culmination of the GLOF Early Warning System project. Following the rigorous testing phase in Week 11, the focus shifted to containerization, orchestration, and cloud deployment. The system is now fully packageable via Docker and deployable to cloud platforms like Render.com, ensuring a resilient and scalable infrastructure for monitoring glacial lake risks.')

    add_heading('2. Containerization (Docker)', 1)
    add_paragraph('To ensure environment parity between development and production, we implemented Dockerfiles for all microservices. The backend is built on a lightweight Python 3.11-slim image, utilizing Gunicorn for high-concurrency request handling.')

    add_paragraph('Code Highlight: Backend Dockerfile', bold=True)
    code_docker = '''FROM python:3.11-slim
WORKDIR /app
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt
COPY . .
EXPOSE 5000
CMD ["gunicorn", "--bind", "0.0.0.0:5000", "--workers", "2", "app:app"]'''
    add_code_block(code_docker)

    add_heading('3. Multi-Service Orchestration', 1)
    add_paragraph('The system architecture consists of the Flask API, a Celery worker for asynchronous notifications, a Redis stream, and the React frontend. We use Docker Compose to orchestrate these services, handling network dependencies and environment variable injection automatically.')

    add_paragraph('Code Highlight: Docker Compose Snippet', bold=True)
    code_compose = '''services:
  backend:
    build: ./backend
    ports: ["5000:5000"]
    env_file: ./backend/.env
  worker:
    build: ./backend
    command: ["python", "-m", "celery", "worker", "-Q", "notifications"]
  frontend:
    build: ./frontend
    ports: ["3000:80"]'''
    add_code_block(code_compose)

    add_heading('4. Cloud Deployment (Render.com)', 1)
    add_paragraph('For production hosting, a "render.yaml" specification was created. This allows for automated "Infrastructure as Code" deployment, connecting the GitHub repository directly to the Render cloud. The configuration handles build commands, static site publishing for the React frontend, and environment management for the Python API.')

    add_paragraph('Code Highlight: Render Infrastructure as Code', bold=True)
    code_render = '''services:
  - type: web
    name: glof-ews-api
    runtime: python
    buildCommand: cd backend && pip install -r requirements.txt
    startCommand: cd backend && bash start.sh
  - type: web
    name: glof-frontend
    runtime: static
    staticPublishPath: frontend/build'''
    add_code_block(code_render)

    add_heading('5. Final Project Conclusion', 1)
    add_paragraph('The GLOF Early Warning System is now complete. Key deliverables achieved include:')
    add_paragraph('• Integrated Telemetry Ingestion with HMAC Security.', bold=True)
    add_paragraph('• Real-time Risk Analysis Engine with SSE Broadcasting.')
    add_paragraph('• Comprehensive Admin Dashboard with Audit Logging.')
    add_paragraph('• Automated CI/CD Pipeline and 70%+ Test Coverage.')
    add_paragraph('• Full Containerization and Cloud-Ready Infrastructure.')
    
    add_paragraph('\\nAll project objectives have been met, and the system is ready for final handover and live monitoring operations.')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Week 12 report successfully generated at: {output_path}")

if __name__ == "__main__":
    update_docx()
