import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def update_docx():
    template_path = r"C:\GLOF\reports\report format.docx"
    output_path = r"C:\GLOF\reports\in_depth\GLOF_Week9_Report.docx"
    
    # Check if template exists
    if not os.path.exists(template_path):
        print(f"Error: Template not found at {template_path}")
        return

    doc = Document(template_path)

    # Find where to start deleting. We'll look for "1. Objectives" or "1. Executive"
    delete_flag = False
    for p in doc.paragraphs:
        if p.text.startswith("1. Objectives") or p.text.startswith("1. Executive"):
            delete_flag = True
        
        if delete_flag:
            # Remove paragraph XML element
            p._element.getparent().remove(p._element)

    def add_heading(text, level):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        return p

    # Injecting our new content for Week 9
    add_heading('1. Executive Summary: Sprint 2', 1)
    add_paragraph('Week 9 (Sprint 2) focused heavily on system integration and test automation. The primary goal was to establish a robust CI/CD pipeline and seamlessly connect the React frontend with the Flask REST APIs and Server-Sent Events (SSE) streams developed in Sprint 1.')

    add_heading('2. CI/CD Integration', 1)
    add_paragraph('To automate testing, linting, and building, a GitHub Actions workflow was introduced. This ensures that any code integrated into the develop or main branches meets our coverage and quality standards.')

    add_heading('2.1 Automated Pipeline (ci.yml)', 2)
    add_paragraph('The pipeline comprises three major jobs: Backend Lint & Test, Frontend Build, and Docker Compose Build. The backend tests are enforced with a minimum coverage threshold of 70%.')

    add_paragraph('Code Highlight: CI/CD Pipeline Configuration', bold=True)
    code_ci = '''name: CI
on:
  push:
    branches: [main, develop]

jobs:
  backend:
    name: Backend — Lint & Test
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - name: Install dependencies
        run: pip install -r requirements.txt flake8
      - name: Run tests with coverage
        run: pytest --cov=core --cov=routes --cov-report=term-missing --cov-fail-under=70

  frontend:
    name: Frontend — Build
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - name: Install dependencies
        run: npm ci
      - name: Build
        run: npm run build'''
    add_code_block(code_ci)

    add_heading('3. Frontend to Backend Integration', 1)
    add_paragraph('Real-time capabilities were integrated into the frontend by consuming the backend SSE (Server-Sent Events) streams. This allows the dashboard to reflect lake telemetry updates instantly without heavy polling.')

    add_heading('3.1 SSE Context Hook (SSEContext.js)', 2)
    add_paragraph('A unified React context was created to manage the EventSource connection, handling reconnect logic automatically and firing native browser notifications upon critical risk detection.')

    add_paragraph('Code Highlight: SSE Stream Consumption', bold=True)
    code_sse = '''const connect = () => {
    const baseUrl = candidates[apiIndexRef.current % candidates.length];
    const eventSource = new EventSource(buildApiUrl(baseUrl, '/api/stream'));
    esRef.current = eventSource;

    eventSource.onmessage = (event) => {
        try {
            const data = JSON.parse(event.data);
            if (data.type === 'connected') return;

            setLatestData(data);
            setLakeMap((prev) => ({ ...prev, [data.lake_id]: data }));
            
            // Notification logic for Critical/High risks
            const isHighRisk = data.risk_level === 'Critical' || data.risk_level === 'High';
            if (data.alert?.alert === true || isHighRisk) {
                fireNativeNotification(data);
            }
        } catch (_) {}
    };

    eventSource.onerror = () => {
        eventSource.close();
        apiIndexRef.current += 1;
        reconnectTimer = window.setTimeout(connect, 4000);
    };
};'''
    add_code_block(code_sse)

    add_heading('4. Test Automation', 1)
    add_paragraph('The backend testing suite was expanded using pytest. Unit tests were written for the core logic components, including the risk engine and the authentication modules, ensuring that both the mathematical models and security flows operate correctly under various scenarios.')

    add_heading('5. Plan for Next Week (Week 10)', 1)
    add_paragraph('Entering Week 10, the project will transition to Phase 5. The focus will be on system optimization, Docker deployment refinement, and stabilizing the production environment on Render.')

    # Make sure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc.save(output_path)
    print(f"Week 9 report successfully generated at: {output_path}")

if __name__ == "__main__":
    update_docx()
