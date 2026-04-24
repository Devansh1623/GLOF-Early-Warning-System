import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def update_docx():
    doc_path = r"C:\GLOF\reports\in_depth\GLOF_Week8_Report.docx"
    doc = Document(doc_path)

    # Find where to start deleting. We'll look for "1. Objectives"
    delete_flag = False
    for p in doc.paragraphs:
        if p.text.startswith("1. Objectives") or p.text.startswith("1. Executive"):
            delete_flag = True
        
        if delete_flag:
            # Remove paragraph XML element
            p._element.getparent().remove(p._element)

    def add_heading(text, level):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        return p

    # Injecting our new content
    add_heading('1. Executive Summary: Sprint 1', 1)
    add_paragraph('Week 8 marks the official transition from system design into active development, kicking off Sprint 1. The primary focus of this sprint was the implementation of core backend features directly derived from the Class and Sequence diagrams established in previous weeks. Furthermore, strict version control practices were instantiated to manage collaboration effectively.')

    add_heading('2. Core Logic Implementation', 1)
    add_paragraph('Following the SOLID principles established in Week 6, the core business logic has been developed into distinct, modular Python packages within the Flask environment.')

    add_heading('2.1 Risk Engine (risk_engine.py)', 2)
    add_paragraph('The mathematical backbone of the system was implemented. It features a hybrid calculation model that utilizes XGBoost and Isolation Forest models for machine learning-based anomaly detection, while maintaining a robust fallback to empirical formulas (incorporating temperature, rainfall, and water level changes based on the Costa 1988 approach).')

    add_paragraph('Code Highlight: Unified Risk Calculator', bold=True)
    code_risk = '''def calculate_risk(temperature, rainfall, water_level_rise, velocity=0.0):
    """
    Unified risk calculator.
    - Uses XGBoost when model is available.
    - Falls back to formula engine otherwise.
    """
    # Build formula breakdown
    t_norm  = _normalize(temperature, TEMP_BASELINE, TEMP_CRITICAL)
    r_norm  = _normalize(rainfall, 0, RAINFALL_CRITICAL)
    wl_norm = _normalize(water_level_rise, 0, WATER_LEVEL_CRITICAL)

    temp_contrib  = round(W_TEMP  * t_norm  * 100, 2)
    rain_contrib  = round(W_RAIN  * r_norm  * 100, 2)
    level_contrib = round(W_LEVEL * wl_norm * 100, 2)
    formula_score = round(min(max(temp_contrib + rain_contrib + level_contrib, 0), 100), 2)

    # Anomaly detection
    anom_pred, anom_raw = anomaly_score(temperature, rainfall, water_level_rise, velocity)

    if _ml_ready:
        try:
            ml_label, ml_conf = _ml_predict(temperature, rainfall, water_level_rise, velocity)
            ml_score = _label_to_score(ml_label)
            # Blend: 60% formula, 40% ML
            blended_score = round(formula_score * 0.60 + ml_score * 0.40, 2)
            level, color  = _level_from_score(blended_score)
            return {"score": blended_score, "level": level, "color": color, "engine": "ml_blended"}
        except Exception as exc:
            pass # fallback

    # Formula fallback
    level, color = _level_from_score(formula_score)
    return {"score": formula_score, "level": level, "color": color, "engine": "formula"}'''
    add_code_block(code_risk)

    add_heading('2.2 Authentication & Security (auth.py)', 2)
    add_paragraph('To secure the REST APIs, Role-Based Access Control (RBAC) was implemented using JSON Web Tokens (JWT). Registration and login controllers were built alongside a password reset flow integrated with the Resend Email API. Strict JSON schema validation using Marshmallow ensures data integrity before database interaction.')

    add_paragraph('Code Highlight: Secure JWT Login Flow', bold=True)
    code_auth = '''@auth_bp.route("/login", methods=["POST"])
def login():
    raw = request.get_json() or {}
    data, errors = validate_json(LoginSchema, raw)
    if errors:
        return jsonify({"error": "Validation failed"}), 422

    email = data["email"].strip().lower()
    user = _db.users.find_one({"email": email})
    
    if not user:
        return jsonify({"error": "Invalid credentials."}), 401

    stored_hash = user["password"]
    if isinstance(stored_hash, str):
        stored_hash = stored_hash.encode("utf-8")

    if not bcrypt.checkpw(data["password"].encode("utf-8"), stored_hash):
        return jsonify({"error": "Invalid credentials."}), 401

    # Generate Secure JWT
    token = create_access_token(
        identity=email,
        additional_claims={"role": user.get("role", "user"), "name": user.get("name", "")},
    )

    return jsonify({"token": token, "user": {"email": email, "role": user.get("role", "user")}}), 200'''
    add_code_block(code_auth)


    add_heading('2.3 Telemetry Simulation (mock_telemetry.py)', 2)
    add_paragraph('To test the pipeline without physical sensors, a simulator was developed. This simulator fetches real-time meteorological data via the Open-Meteo API and calculates diurnal temperature drifts and randomized water level spikes to mimic genuine environmental behaviors.')

    add_paragraph('Code Highlight: Blending Real Weather with Simulation Spikes', bold=True)
    code_sim = '''def simulate_reading(lake, tick):
    lid = lake["id"]
    state = lake_state[lid]

    # Fetch real weather via Open-Meteo
    real = fetch_open_meteo(lake["lat"], lake["lon"])

    # Temperature: blend real + diurnal cycle + noise
    hour = datetime.utcnow().hour
    diurnal = 6 * math.sin((hour - 6) / 24 * 2 * math.pi)
    if real["temperature"] is not None:
        temp = real["temperature"] + diurnal * 0.3 + random.gauss(0, 1.5)
    else:
        temp = 5 + random.gauss(0, 3) + diurnal + random.gauss(0, 1)

    # Water level: gradual drift up + occasional spike
    state["water_level"] += state["drift"] + random.gauss(0, 2)
    if random.random() < 0.03:
        state["water_level"] += random.uniform(20, 60)
        
    if state["water_level"] > 350:
        state["water_level"] = random.uniform(20, 80)
        state["drift"] = random.uniform(0.1, 0.5)

    return {
        "lake_id": lid,
        "temperature": round(temp, 2),
        "water_level_rise": round(max(0, state["water_level"]), 2),
        "timestamp": datetime.utcnow().isoformat() + "Z",
    }'''
    add_code_block(code_sim)


    add_heading('3. Git Branching Strategy', 1)
    add_paragraph('To maintain a clean and reliable codebase as complexity increases, a formalized Git workflow has been strictly enforced across the team.')
    add_paragraph('Feature Branching Workflow:', bold=True)
    add_paragraph('• Main Branch (main): Represents the production-ready state. Direct commits to this branch are strictly prohibited.')
    add_paragraph('• Develop Branch (develop): Serves as the primary integration branch for the next release.')
    add_paragraph('• Feature Branches: Individual tasks (e.g., feat/risk-engine) are branched off develop. Once complete, a Pull Request is opened to merge back into develop.')

    add_heading('4. Plan for Next Week (Week 9)', 1)
    add_paragraph('Entering Week 9, the project will advance into Sprint 2. The primary objectives include:')
    add_paragraph('• Integration of Data Sources: Connecting the frontend React UI hooks deeply with the established Flask REST endpoints and Server-Sent Events (SSE) streams.')
    add_paragraph('• Development of Automated Pipelines: Setting up GitHub Actions or similar automated triggers for Continuous Integration testing.')

    doc.save(doc_path)
    print("Done!")

if __name__ == "__main__":
    update_docx()
