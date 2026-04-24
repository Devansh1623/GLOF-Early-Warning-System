import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_week1_report():
    doc = Document()
    
    # Title
    title = doc.add_heading("Week 1 Project Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run("Project Title: ").bold = True
    p.add_run("GLOF Early Warning System (GLOFWatch)\n")
    p.add_run("Phase 1: ").bold = True
    p.add_run("Planning & Requirement Analysis\n")
    p.add_run("Focus: ").bold = True
    p.add_run("Defining the scope, setting up infrastructure, and identifying functional requirements.\n")
    p.add_run("Date Range: ").bold = True
    p.add_run("19 - 24 January 2026\n")
    
    # 1. Executive Summary & Problem Statement
    doc.add_heading("1. Executive Summary & Problem Statement", level=1)
    doc.add_paragraph(
        "During Week 1, the core project team was successfully formed, and the primary focus was on establishing "
        "a robust project foundation. The team engaged in extensive problem definition exercises to map out the "
        "impact of Glacial Lake Outburst Floods (GLOFs) in the Himalayan region."
    )
    doc.add_heading("Problem Statement:", level=2)
    doc.add_paragraph(
        "Due to accelerated climate change, high-altitude glacial lakes are expanding rapidly. The existing monitoring "
        "infrastructure is sparse and often lacks real-time predictive capabilities, leaving downstream communities "
        "vulnerable to catastrophic sudden floods. Our proposed solution is a real-time, Machine Learning-enhanced "
        "Early Warning System (EWS) that integrates simulated IoT sensor telemetry, Open-Meteo weather data, and rigorous "
        "predictive risk algorithms to deliver actionable alerts with minimal latency."
    )
    
    # 2. Tooling and Infrastructure Setup
    doc.add_heading("2. Tooling and Infrastructure Setup", level=1)
    doc.add_paragraph(
        "Following strict engineering practices, the initial development environment and tracking systems were configured:"
    )
    
    doc.add_heading("Git Repository Initialization", level=3)
    p_git = doc.add_paragraph(style='List Bullet')
    p_git.add_run("Platform: ").bold = True
    p_git.add_run("GitHub was selected for version control.\n")
    p_git = doc.add_paragraph(style='List Bullet')
    p_git.add_run("Repository: ").bold = True
    p_git.add_run("Devansh1623/GLOF-Early-Warning-System\n")
    p_git = doc.add_paragraph(style='List Bullet')
    p_git.add_run("Configuration: ").bold = True
    p_git.add_run("Configured branching strategies protecting the 'main' and 'develop' branches. "
                  "A comprehensive .gitignore was established covering Python (Flask) and Node.js (React) environments.")
    
    doc.add_heading("Agile Project Management (Jira Setup)", level=3)
    p_jira = doc.add_paragraph(style='List Bullet')
    p_jira.add_run("Tool: ").bold = True
    p_jira.add_run("Jira Software.\n")
    p_jira = doc.add_paragraph(style='List Bullet')
    p_jira.add_run("Methodology: ").bold = True
    p_jira.add_run("A Scrum board was configured to track progress over the upcoming 10-week cycle.\n")
    p_jira = doc.add_paragraph(style='List Bullet')
    p_jira.add_run("Epics Created: ").bold = True
    p_jira.add_run("Data Ingestion Pipeline, ML Risk Engine, Real-time Alerting, and Frontend UI Dashboard.")
    
    # 3. Deliverable: Project Vision Document
    doc.add_heading("3. Deliverable: Project Vision Document", level=1)
    doc.add_paragraph(
        "A formal Project Vision Document was drafted to guide architectural decisions. The system is envisioned as "
        "a decoupled, production-grade application:"
    )
    p_vision = doc.add_paragraph(style='List Bullet')
    p_vision.add_run("Backend: ").bold = True
    p_vision.add_run("Python Flask providing RESTful APIs and SSE (Server-Sent Events) for real-time streaming.")
    p_vision = doc.add_paragraph(style='List Bullet')
    p_vision.add_run("Frontend: ").bold = True
    p_vision.add_run("React.js leveraging 'Glacial Observatory' Design tokens for a high-fidelity visual experience.")
    p_vision = doc.add_paragraph(style='List Bullet')
    p_vision.add_run("Data Storage: ").bold = True
    p_vision.add_run("MongoDB for persistent storage (Lakes, Telemetry, Alerts) and Redis for fast in-memory caching and Pub/Sub.")
    p_vision = doc.add_paragraph(style='List Bullet')
    p_vision.add_run("Intelligence: ").bold = True
    p_vision.add_run("Hybrid Risk Engine blending empirical formulas (Costa equations) with an XGBoost/Isolation Forest model.")

    # 4. Deliverable: Initial Product Backlog
    doc.add_heading("4. Deliverable: Initial Product Backlog", level=1)
    doc.add_paragraph(
        "An initial Product Backlog was created in Jira, breaking down the Epics into actionable user stories. "
        "High-priority stories prioritized for early sprints include:"
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Epic'
    hdr_cells[1].text = 'User Story / Task'
    hdr_cells[2].text = 'Priority'
    
    backlog_items = [
        ("Data Ingestion", "As a sensor, I need a secure REST API to POST telemetry data (temperature, rainfall, water level).", "P0"),
        ("Data Ingestion", "As an Admin, I need the database seeded with 12 real CWC-monitored glacial lakes.", "P0"),
        ("Risk Engine", "As the system, I need an algorithm to calculate risk scores based on incoming telemetry.", "P0"),
        ("Auth & Security", "As a User, I need JWT-based authentication to securely access the dashboard.", "P1"),
        ("Alerting", "As an Admin, I need real-time alerts pushed to the UI via Server-Sent Events (SSE).", "P1"),
        ("Frontend UI", "As an Admin, I need a map view (Leaflet) displaying all monitored glacial lakes.", "P1")
    ]
    
    for epic, story, priority in backlog_items:
        row_cells = table.add_row().cells
        row_cells[0].text = epic
        row_cells[1].text = story
        row_cells[2].text = priority

    # 5. Goals for Next Week
    doc.add_heading("5. Goals for Next Week (Week 2)", level=1)
    doc.add_paragraph("With the project infrastructure established, the goals for Week 2 are:", style='List Bullet')
    doc.add_paragraph("Identify all System Actors (IoT Sensors, Admins, Researchers).", style='List Bullet')
    doc.add_paragraph("Map out all functional requirements through Use Cases.", style='List Bullet')
    doc.add_paragraph("Create the Use Case Diagram and detailed Use Case Descriptions.", style='List Bullet')

    p_footer = doc.add_paragraph("\n\nMentor Signature: __________________      Project Guide Signature: __________________")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(r"c:\GLOF\reports\in_depth", exist_ok=True)
    filename = r"c:\GLOF\reports\in_depth\GLOF_Week_1_Report_InDepth.docx"
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    create_week1_report()
