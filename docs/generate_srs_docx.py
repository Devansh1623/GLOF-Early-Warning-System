import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def generate_srs():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Software Requirements Specification (SRS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('GLOFWatch - Glacial Lake Outburst Flood Early Warning System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)
    
    doc.add_paragraph('\nVersion: 1.0\nDate: April 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    
    add_heading(doc, '1.1 Purpose', 2)
    add_paragraph(doc, 'The purpose of this document is to outline the software requirements for GLOFWatch, an open-source, ML-augmented Early Warning System (EWS) designed to mitigate the risks of Glacial Lake Outburst Floods (GLOFs) in the Indian Himalayan Region. This specification dictates the system architecture, feature sets, technical constraints, semantic architectures, and operating environments.')
    
    add_heading(doc, '1.2 Intended Audience', 2)
    add_paragraph(doc, 'This document is intended for Disaster Management Authorities (e.g., NDMA, CWC) to understand operational capabilities; Software Engineers & Researchers for technical onboarding to the codebase, frontend mechanics, and ML architecture; and Academic Reviewers to evaluate the system\'s scope, adherence to disaster risk reduction standards, and SEO compliance.')
    
    add_heading(doc, '1.3 Product Scope', 2)
    add_paragraph(doc, 'GLOFWatch is a full-stack intelligence platform that integrates hardware telemetry processing, real-time Machine Learning (ML) inference, spatial mapping, and critical alert dispatching. Unlike fragmented manual observation systems, GLOFWatch acts as a continuous automated operational pipeline. It ingests environmental data (water levels, seismic activity, temperature, precipitation) from lake-side sensors every 5 seconds, calculates real-time risk severity using both legacy formulas (e.g., Costa, 1988) and an XGBoost/IsolationForest machine learning backend, and instantly dispatches warnings to field commanders and public officials through a unified GIS dashboard and email/SMS alerts.')
    
    add_heading(doc, '1.4 References', 2)
    add_paragraph(doc, '• Sendai Framework for Disaster Risk Reduction 2015-2030\n• NDMA Guidelines on Management of Glacial Lake Outburst Floods\n• UN Sustainable Development Goals (SDG 11, SDG 13)\n• W3C Web Content Accessibility Guidelines (WCAG) 2.1')
    
    doc.add_page_break()

    # 2. Overall Description & Diagrams
    add_heading(doc, '2. Overall Description', 1)
    
    add_heading(doc, '2.1 Product Perspective', 2)
    add_paragraph(doc, 'GLOFWatch operates as an independent, cloud-hosted web application comprised of separated microservices. It features a React-based frontend built heavily for GIS visualization and Search Engine Optimization (SEO), and a Python/Flask-based backend built for high-throughput concurrency.')
    
    add_heading(doc, '2.2 System Architecture Diagrams', 2)
    add_paragraph(doc, 'The following architectural diagrams visually describe the object-oriented structure, microservice packaging, and physical deployment boundaries of the GLOFWatch system.')
    
    report_dir = r"c:\GLOF\reports\in_depth"
    
    # Deployment Diagram
    dep_path = os.path.join(report_dir, "deployment_diagram.png")
    if os.path.exists(dep_path):
        add_paragraph(doc, 'Deployment Architecture', bold=True)
        doc.add_picture(dep_path, width=Inches(6.0))
        doc.add_paragraph('Figure 1: Final Production Deployment Diagram (Render, Redis, MongoDB Atlas)')
    
    # Class Diagram
    class_path = os.path.join(report_dir, "detailed_class_diagram.png")
    if os.path.exists(class_path):
        add_paragraph(doc, '\nClass Data Model', bold=True)
        doc.add_picture(class_path, width=Inches(6.0))
        doc.add_paragraph('Figure 2: Detailed Software Class Definitions')
        
    doc.add_page_break()

    # 3. System Features
    add_heading(doc, '3. System Features', 1)
    
    add_heading(doc, '3.1 Telemetry Ingestion Pipeline', 2)
    add_paragraph(doc, 'A RESTful POST endpoint enabling hardware sensors (or the telemetry simulator) to push JSON packet updates containing water_level, temperature, precipitation_mm, and seismic_activity_g. The API MUST be capable of processing a minimum of 1,000 asynchronous concurrent connections. Payloads MUST be validated against database schemas before database insertion.')
    
    add_heading(doc, '3.2 Real-time Risk Inference Engine', 2)
    add_paragraph(doc, 'The calculation core. Invoked on every new telemetry packet. It MUST combine the Costa (1988) empirical formula (weighted 60%) with a pre-trained XGBoost Model + Isolation Forest (weighted 40%). Execution must occur within <100ms, returning a risk severity score normalized between 0.00 and 100.00.')
    
    add_heading(doc, '3.3 Event-Driven Alerting (SSE & Queues)', 2)
    add_paragraph(doc, 'Broadcast mechanisms to inform users immediately when thresholds are crossed. Warnings (Score > 60) must generate an internal audit log and SSE UI notification. Emergencies (Score > 80) must dispatch background priority tasks sending emails via Resend API to all subscribed operators.')
    
    add_heading(doc, '3.4 Interactive Geospatial Display', 2)
    add_paragraph(doc, 'A GIS Map UI displaying all monitored lakes. Rendered using React-Leaflet. Active lakes MUST render as interactive marker pins whose colors dynamically change based on real-time risk scores without requiring page reloads.')
    
    doc.add_page_break()

    # 4. Web Accessibility and SEO Requirements
    add_heading(doc, '4. SEO and Accessibility Specifications', 1)
    add_paragraph(doc, 'To ensure the platform is highly discoverable by researchers, government agencies, and indexed correctly by search engines, the Frontend React Application strictly adheres to modern technical SEO and WCAG usability standards.')
    
    add_heading(doc, '4.1 Technical SEO Architecture', 2)
    add_paragraph(doc, '• React Helmet Integration: All pages MUST utilize `react-helmet` to inject specific metadata, title tags, and localized descriptions per geographical zone.\n• Canonical URLs: Dynamic pages must define canonical URLs to prevent duplicate content indexing by Googlebot.\n• Semantic HTML5 structure: The layout must rigidly follow correct tag nesting (one `<h1>` per page, hierarchical `<h2-h6>` tags, `<header>`, `<main>`, `<article>`, `<nav>`, `<footer>`).\n• Core Web Vitals: Largest Contentful Paint (LCP) must load within 2.5 seconds; Cumulative Layout Shift (CLS) must be < 0.1.')
    
    add_heading(doc, '4.2 Accessibility (WCAG 2.1 AA)', 2)
    add_paragraph(doc, '• ARIA Labels: All interactive elements, specifically the interactive Leaflet Map pins, MUST contain `aria-labels` and `roles` allowing screen readers to interpret glacial risk states accurately.\n• Color Contrast: The dark navy base aesthetic MUST pass the 4.5:1 contrast ratio against text identifiers ensuring viability for visually impaired analysts.\n• Keyboard Navigation: The entire operations dashboard MUST be fully navigable using `Tab` indexing without trapping focus in modals.')
    
    # 5. Non-Functional Requirements
    add_heading(doc, '5. Non-Functional Requirements', 1)
    add_heading(doc, '5.1 Performance', 2)
    add_paragraph(doc, 'Telemetry ingestion to visual dashboard graph update (end-to-end latency) MUST be under exactly 5.0 seconds. API endpoints MUST respond in < 250 milliseconds at the 95th percentile under standard load.')
    
    add_heading(doc, '5.2 Security and Compliance', 2)
    add_paragraph(doc, 'All REST transit MUST be strictly over TLS 1.3 (HTTPS). JWT tokens must have hardened expiration policies. Passwords are mathematically hashed natively using bcrypt and explicitly scrubbed from any API JSON responses.')

    output_dir = r"c:\GLOF\docs"
    file_path = os.path.join(output_dir, "GLOF_Comprehensive_SRS.docx")
    
    try:
        doc.save(file_path)
        print(f"Successfully generated {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    generate_srs()
