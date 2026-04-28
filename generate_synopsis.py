import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    return p

def create_synopsis():
    doc = Document()
    
    # Title Page
    title = doc.add_paragraph("PROJECT SYNOPSIS\n")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    logo_path = 'c:/GLOF/svvv_logo.png'
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_r = logo_p.add_run()
        logo_r.add_picture(logo_path, width=Inches(1.5))
    
    proj_title = doc.add_paragraph("\nGLOFWatch — Glacial Lake Outburst Flood Early Warning System\n")
    proj_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_run = proj_title.runs[0]
    pt_run.font.name = 'Arial'
    pt_run.font.size = Pt(22)
    pt_run.font.bold = True
    pt_run.font.color.rgb = RGBColor(0, 51, 102)
    
    sub = doc.add_paragraph("Submitted in partial fulfillment of the requirements for the award of the degree of\n")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.name = 'Arial'
    
    degree = doc.add_paragraph("Bachelor of Technology (B.Tech) - Minor Project\n")
    degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
    degree.runs[0].font.bold = True
    degree.runs[0].font.size = Pt(14)
    degree.runs[0].font.name = 'Arial'
    
    branch = doc.add_paragraph("Computer Science and Engineering (Artificial Intelligence in Collaboration with IBM)\n3rd Year, 6th Semester\n")
    branch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in branch.runs:
        r.font.name = 'Arial'
    
    doc.add_paragraph("\n")
    
    team_table = doc.add_table(rows=1, cols=2)
    team_table.style = 'Table Grid'
    
    hdr_cells = team_table.rows[0].cells
    p1 = hdr_cells[0].paragraphs[0]
    p1.add_run('Submitted To:').bold = True
    p2 = hdr_cells[1].paragraphs[0]
    p2.add_run('Submitted By:').bold = True
    
    row_cells = team_table.add_row().cells
    row_cells[0].text = "Prof. Jagrati Nagdiya (Guide)\nProf. Minakshi Jalandra (Mentor)"
    row_cells[1].text = "Devansh Geria (Leader)\nEnrollment: 23100BTCSAII14280\n\nShagun Soni\nEnrollment: 23100BTCSAII14322\n\nTanay Puranik\nEnrollment: 23100BTCSAII14327\n\nVashishth Kasare\nEnrollment: 23100BTCSAII14333"
    
    doc.add_paragraph("\n")
    
    inst = doc.add_paragraph("Shri Vaishnav Institute of Information and Technology (SVITS)\nShri Vaishnav Vidyapeeth Vishwavidyalaya (SVVV)\nIndore, Madhya Pradesh")
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst.runs[0].font.bold = True
    inst.runs[0].font.size = Pt(14)
    inst.runs[0].font.name = 'Arial'
    
    doc.add_page_break()
    
    # 1. Introduction
    add_heading(doc, "1. Introduction", 1)
    add_paragraph(doc, "Glacial Lake Outburst Floods (GLOFs) are among the most catastrophic natural hazards in high-mountain regions, capable of releasing millions of cubic meters of water within minutes. As climate change accelerates glacial retreat across the Hindu Kush–Himalayan belt, the number and size of glacial lakes is growing rapidly, intensifying downstream risk for communities, infrastructure, and hydroelectric assets. Currently, the Indian Himalayan Region contains over 2,400 glacial lakes, with 188 classified as potentially dangerous by NRSC/ISRO.")
    add_paragraph(doc, "GLOFWatch is an open-source, ML-augmented Early Warning System (EWS) designed to mitigate these risks. It operates as an independent, cloud-hosted platform that integrates hardware telemetry processing, real-time Machine Learning (ML) inference, spatial mapping, and critical alert dispatching. Unlike fragmented manual observation systems, GLOFWatch acts as a continuous automated operational pipeline, ingesting environmental data from lake-side sensors every 5 seconds.")
    
    # 2. Problem Statement
    add_heading(doc, "2. Problem Statement", 1)
    add_paragraph(doc, "Despite the significant threat posed by GLOFs—exemplified by tragic events in Kedarnath, Chamoli, and South Lhonak—current monitoring systems in the Indian Himalayan Region suffer from several critical shortcomings:")
    add_paragraph(doc, "• Sparse Sensor Coverage: Relying heavily on manual, seasonal field surveys rather than continuous automated telemetry.")
    add_paragraph(doc, "• Delayed Alerting: A lack of real-time processing results in hours-long lag between an event initiation and the dissemination of alerts.")
    add_paragraph(doc, "• Siloed Data Sources: Weather, water level, and geological data exist in separate systems, making holistic risk assessment difficult.")
    add_paragraph(doc, "• Binary Risk Assessment: Current approaches often use simple alert/no-alert thresholds rather than nuanced risk scoring based on compounded environmental variables.")
    add_paragraph(doc, "GLOFWatch addresses these gaps by providing an integrated, automated, always-on monitoring dashboard with hybrid mathematical and machine learning-based risk inference.")
    
    # 3. Objectives
    add_heading(doc, "3. Objectives", 1)
    add_paragraph(doc, "1. Real-Time Telemetry Ingestion: Continuously aggregate real-time weather and sensor telemetry (water level, temperature, rainfall, velocity) for Himalayan glacial lakes.")
    add_paragraph(doc, "2. Hybrid Risk Scoring Engine: Implement a robust scoring engine combining a weighted empirical formula (based on Costa, 1988) with an XGBoost Classifier for predictive insights.")
    add_paragraph(doc, "3. Anomaly Detection: Utilize an Isolation Forest machine learning model to detect faulty or anomalous sensor readings to prevent false positives.")
    add_paragraph(doc, "4. Multi-Channel Alerting: Provide instant warnings (Server-Sent Events, toast notifications, email via Resend API) when critical thresholds are breached.")
    add_paragraph(doc, "5. Geospatial Visualization: Build an interactive map dashboard using Leaflet to dynamically plot lake basins with their real-time severity statuses.")
    
    # 4. Scope and Methodology
    add_heading(doc, "4. Scope & Methodology", 1)
    add_heading(doc, "4.1 System Scope", 2)
    add_paragraph(doc, "The system encompasses the end-to-end pipeline required for disaster mitigation: from simulated IoT data ingestion (with Open-Meteo weather integration) to risk computation, spatial visualization, alert dispatching, and audit logging. The platform is designed to scale horizontally across hundreds of lakes.")
    
    add_heading(doc, "4.2 Methodology", 2)
    add_paragraph(doc, "• Data Ingestion & Simulation: The simulator generates realistic telemetry blending real weather from Open-Meteo with synthetic diurnal cycles and stochastic events (e.g., sudden water level spikes). Data is pushed to the Flask backend every 5 seconds.")
    add_paragraph(doc, "• Hybrid Risk Engine: Upon receiving telemetry, the system calculates a base score using an empirical formula evaluating Water Level Rise (35%), Temperature (35%), and Rainfall (30%). Concurrently, an XGBoost model predicts the risk class. The final risk score is a 60/40 blend of the formula and ML scores.")
    add_paragraph(doc, "• Alerting Mechanism: The engine evaluates the risk against defined thresholds (>60 High, >80 Critical). Alerts are pushed in real-time to connected web clients via Server-Sent Events (SSE) and critical alerts trigger asynchronous email dispatches using Celery to prevent UI blocking.")
    
    doc.add_page_break()
    
    # 5. System Architecture & Diagrams
    add_heading(doc, "5. System Architecture & Design", 1)
    add_paragraph(doc, "The architecture of GLOFWatch follows a microservices pattern, decoupling the frontend GIS application, the backend intelligence API, the background task workers, and the telemetry simulators.")
    
    # Insert Component Diagram if exists
    comp_path = 'c:/GLOF/reports/in_depth/component_diagram.png'
    if os.path.exists(comp_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(comp_path, width=Inches(6))
        p_cap = doc.add_paragraph("Figure 1: System Component Diagram")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_paragraph(doc, "The system leverages Docker Compose for local orchestration, launching four primary containers: the React frontend (served via Nginx), the Flask backend API (served via Gunicorn), the Celery worker for background jobs, and the Mock Telemetry generator.")
    
    # Insert Deployment Diagram if exists
    dep_path = 'c:/GLOF/reports/in_depth/deployment_diagram.png'
    if os.path.exists(dep_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(dep_path, width=Inches(6))
        p_cap = doc.add_paragraph("Figure 2: Deployment Diagram (Render Cloud)")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    
    # Insert Class Diagram if exists
    class_path = 'c:/GLOF/reports/in_depth/detailed_class_diagram.png'
    if os.path.exists(class_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(class_path, width=Inches(6))
        p_cap = doc.add_paragraph("Figure 3: Detailed Class Diagram")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 6. Technology Stack
    add_heading(doc, "6. Technology Stack", 1)
    add_paragraph(doc, "The project relies on a modern, scalable open-source stack:")
    add_paragraph(doc, "• Frontend: React 18, React Router DOM, Axios, React-Leaflet, Recharts, Tailwind CSS. Providing a responsive, component-based GIS UI.")
    add_paragraph(doc, "• Backend: Python 3.11, Flask 3.1, Flask-JWT-Extended, Celery, Gunicorn. Ensuring high concurrency and asynchronous task processing.")
    add_paragraph(doc, "• Machine Learning: Scikit-learn (Isolation Forest for anomaly detection), XGBoost (Classification), Pandas, NumPy.")
    add_paragraph(doc, "• Database & Cache: MongoDB Atlas (Document Store for persistence), Upstash Redis (In-memory caching and Pub/Sub for SSE messaging).")
    add_paragraph(doc, "• Infrastructure: Docker, GitHub Actions (CI/CD), Render.com for cloud hosting.")
    
    # 7. Project Duration & Timeline
    add_heading(doc, "7. Project Duration", 1)
    add_paragraph(doc, "The project was executed over a 12-week development lifecycle encompassing:")
    add_paragraph(doc, "• Weeks 1-2: Requirement gathering, architectural scaffolding, database schema design, and secure authentication (JWT).")
    add_paragraph(doc, "• Weeks 3-5: Risk engine implementation, telemetry pipeline simulation, interactive map integration, and Server-Sent Events (SSE) streaming.")
    add_paragraph(doc, "• Weeks 6-8: Machine Learning model training (12K synthetic samples), Alert notification system (Celery + Resend), and Admin dashboard execution.")
    add_paragraph(doc, "• Weeks 9-12: System stabilization, unit testing (Pytest >70% coverage), CI/CD pipeline automation, and production deployment on Render.com.")
    
    # 8. Future Scope
    add_heading(doc, "8. Future Scope", 1)
    add_paragraph(doc, "The foundational architecture of GLOFWatch permits significant future enhancements:")
    add_paragraph(doc, "1. Real IoT Sensor Integration: Substituting the simulator with physical water pressure sensors and weather stations located at high-altitude moraines.")
    add_paragraph(doc, "2. Satellite Imagery Analysis: Utilizing Sentinel-2 optical imagery and SAR data for automated Lake Area extent changes (NDWI).")
    add_paragraph(doc, "3. Deep Learning Forecasting: Implementing LSTM or Transformer neural networks for predictive time-series forecasting of lake volume increases over weeks/months.")
    add_paragraph(doc, "4. Enhanced Dissemination: Integration with SMS Gateways (Twilio) to reach rural Himalayan populations lacking continuous internet connectivity.")
    
    doc.save("c:/GLOF/reports/GLOF_Project_Synopsis_Official.docx")
    print("Synopsis generated successfully at c:/GLOF/reports/GLOF_Project_Synopsis_Official.docx")

if __name__ == '__main__':
    create_synopsis()
